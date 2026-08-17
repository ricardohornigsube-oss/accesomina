from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path('/Users/ricardo.hornig/Desktop/Nexo Klar/Oferta_Comercial_Unificada_Nexo_Klar.docx')
REPO_COPY = Path('/Users/ricardo.hornig/Desktop/Acceso Mina /Acceso Mina/docs/Oferta_Comercial_Unificada_Nexo_Klar.docx')
INDIGO = '312E81'; MAGENTA = 'DB2777'; TEAL = '0F766E'; INK = '121826'; MUTED = '536177'; PALE = 'F7F7FB'; LINE = 'D9DDEA'; WHITE = 'FFFFFF'

def cell_shade(cell, color):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), color); tcPr.append(shd)

def cell_border(cell, color=LINE):
    tcPr = cell._tc.get_or_add_tcPr(); borders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right'):
        e = OxmlElement(f'w:{edge}'); e.set(qn('w:val'),'single'); e.set(qn('w:sz'),'8'); e.set(qn('w:color'),color); borders.append(e)
    tcPr.append(borders)

def set_cell_text(cell, text, bold=False, color=INK, size=9.2):
    cell.text = ''
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)
    r = p.add_run(str(text)); r.bold = bold; r.font.name='Arial'; r.font.size=Pt(size); r.font.color.rgb=RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def set_repeat_header(row):
    trPr = row._tr.get_or_add_trPr(); el=OxmlElement('w:tblHeader'); el.set(qn('w:val'),'true'); trPr.append(el)

def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'
    hdr=t.rows[0]; set_repeat_header(hdr)
    for i,h in enumerate(headers):
        if widths: hdr.cells[i].width = Inches(widths[i])
        cell_shade(hdr.cells[i], INDIGO); cell_border(hdr.cells[i], INDIGO); set_cell_text(hdr.cells[i],h,True,WHITE,9)
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            if widths: cells[i].width=Inches(widths[i])
            cell_shade(cells[i], WHITE if len(t.rows)%2 else PALE); cell_border(cells[i]); set_cell_text(cells[i],v,False,INK,8.8)
    doc.add_paragraph().paragraph_format.space_after=Pt(3)
    return t

def title(doc, text, sub=None):
    p=doc.add_paragraph(); p.style='Heading 1'; p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(4)
    r=p.add_run(text); r.font.name='Arial'; r.font.size=Pt(24); r.bold=True; r.font.color.rgb=RGBColor.from_string(INK)
    if sub:
        p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(10); r=p.add_run(sub); r.font.name='Arial'; r.font.size=Pt(11.5); r.font.color.rgb=RGBColor.from_string(MUTED)

def h2(doc, text, sub=None):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(3)
    r=p.add_run(text); r.font.name='Arial'; r.font.size=Pt(16); r.bold=True; r.font.color.rgb=RGBColor.from_string(INDIGO)
    if sub:
        p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(7); r=p.add_run(sub); r.font.name='Arial'; r.font.size=Pt(10.5); r.font.color.rgb=RGBColor.from_string(MUTED)

def body(doc, text, bold_prefix=None):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(6); p.paragraph_format.line_spacing=1.15
    if bold_prefix:
        r=p.add_run(bold_prefix); r.bold=True; r.font.name='Arial'; r.font.size=Pt(10.2); r.font.color.rgb=RGBColor.from_string(INK)
    r=p.add_run(text); r.font.name='Arial'; r.font.size=Pt(10.2); r.font.color.rgb=RGBColor.from_string(INK)

def bullets(doc, items):
    for item in items:
        p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after=Pt(2); p.paragraph_format.line_spacing=1.05
        r=p.add_run(item); r.font.name='Arial'; r.font.size=Pt(9.7); r.font.color.rgb=RGBColor.from_string(INK)

def callout(doc, label, text):
    t=doc.add_table(rows=1, cols=2); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    t.columns[0].width=Inches(1.25); t.columns[1].width=Inches(5.75)
    for i,content in enumerate([label,text]):
        c=t.cell(0,i); c.width=Inches(1.25 if i==0 else 5.75); cell_border(c,'9BE5DC'); cell_shade(c,'EAFBF9'); set_cell_text(c,content,i==0,TEAL if i==0 else INK,9.3)
    doc.add_paragraph().paragraph_format.space_after=Pt(3)

def page(doc): doc.add_page_break()

def add_header_footer(section):
    header=section.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    r=header.add_run('NEXO KLAR SPA · OFERTA SaaS 2026–2027'); r.font.name='Arial'; r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string(MUTED)
    footer=section.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=footer.add_run('Documento comercial interno · Valores netos referenciales en CLP · Agosto 2026'); r.font.name='Arial'; r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string(MUTED)

doc=Document()
sec=doc.sections[0]; sec.top_margin=Inches(.62); sec.bottom_margin=Inches(.6); sec.left_margin=Inches(.68); sec.right_margin=Inches(.68); add_header_footer(sec)
styles=doc.styles; styles['Normal'].font.name='Arial'; styles['Normal'].font.size=Pt(10)

# Cover
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(58); r=p.add_run('NEXO KLAR SPA'); r.font.name='Arial'; r.font.size=Pt(13); r.bold=True; r.font.color.rgb=RGBColor.from_string(TEAL)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(12); r=p.add_run('Oferta comercial SaaS\nunificada'); r.font.name='Arial'; r.font.size=Pt(35); r.bold=True; r.font.color.rgb=RGBColor.from_string(INK)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(24); r=p.add_run('Dos líneas de servicio, seis opciones de capacidad y una ruta clara desde las personas hasta el control total de la operación.'); r.font.name='Arial'; r.font.size=Pt(15); r.font.color.rgb=RGBColor.from_string(MUTED)
callout(doc,'PROPUESTA','Nexo Klar ordena personas, cumplimiento, órdenes, recursos y evidencia en una misma plataforma. El cliente escala por necesidad real, sin perder información ni historial.')
h2(doc,'Cómo leer esta propuesta')
table(doc,['Línea','Qué resuelve','Para quién'],[
['Nexo Klar Inicial','Controla personas, documentación, clientes, contratos y órdenes de servicio.','Empresas que buscan orden, cumplimiento y visibilidad operacional.'],
['Nexo Klar Platinum','Añade coordinación de recursos, terceros, activos, inventario, costos y analítica.','Empresas que operan múltiples servicios, equipos y recursos en terreno.'],
['Adicionales independientes','Libro de Obra y Gestor Comercial.','Se activan en cualquiera de las seis opciones sin cambiar la base de datos.']
],[1.55,3.2,2.25])
body(doc,'Los valores son referenciales netos en CLP. Se ajustan mediante una cotización por industria, nivel de implementación, almacenamiento, consumo de mensajería, OCR, firma electrónica e integraciones.')
page(doc)

title(doc,'1. Principio comercial','Una oferta simple de explicar y fácil de escalar')
body(doc,'La propuesta reemplaza una estructura de tres planes genéricos por dos líneas muy claras. La primera permite ordenar y cumplir; la segunda permite operar, medir y escalar. Dentro de cada línea, la capacidad aumenta sin obligar al cliente a rediseñar su forma de trabajar.')
table(doc,['Regla','Aplicación comercial'],[
['Cobro por escala y valor','La tarifa considera personas activas, usuarios nominados, órdenes activas, capacidad de almacenamiento y nivel de soporte; no el número de pantallas que el cliente abre.'],
['Misma base, mayor profundidad','Un cliente puede subir de opción o de línea sin perder su historial, usuarios, documentos o configuraciones.'],
['Servicios adicionales independientes','Libro de Obra y Gestor Comercial se venden por separado para no cargar el plan base de empresas que aún no los necesitan.'],
['Implementación visible','Carga, depuración, parametrización e integraciones se cotizan aparte: son trabajo especializado y no deben quedar ocultos en la mensualidad.']
],[2.1,4.9])
callout(doc,'Referente de mercado','Las plataformas de field service suelen combinar niveles funcionales, capacidad de uso, acompañamiento y complementos. Por ejemplo, Housecall Pro diferencia planes por usuarios y soporte; ServiceM8 utiliza capacidad de trabajos y mensajería; FieldPulse usa precios por técnico. Nexo Klar adapta esa lógica al contexto operacional chileno.')
h2(doc,'Cliente ideal')
bullets(doc,[
'Empresas de servicios, mantenimiento, construcción, energía, logística, facilities, seguridad, agroindustria y minería con 20 a 500 personas.',
'Operaciones que dependen de planillas, carpetas, correos y personas clave para saber quién puede trabajar y con qué recursos.',
'Equipos que necesitan demostrar cumplimiento documental, coordinar servicios y mantener historial verificable frente a clientes, mandantes y auditorías.'
])
page(doc)

title(doc,'2. Nexo Klar Inicial','Orden y cumplimiento que se transforma en control diario')
body(doc,'La línea Inicial está diseñada para instalar una fuente única de información. Parte por las personas y sus requisitos; luego las vincula con clientes, contratos y órdenes de servicio para saber qué está listo, pendiente o restringido.')
table(doc,['Opción','Inicio','Equipo','Empresa'],[
['Precio mensual neto CLP','$169.000','$289.000','$429.000'],
['Personas activas incluidas','Hasta 20','Hasta 60','Hasta 120'],
['Usuarios nominados','3','8','15'],
['Órdenes de servicio activas/mes','10','40','100'],
['Almacenamiento incluido','10 GB','35 GB','100 GB'],
['Soporte','Correo hábil','Correo + sesión mensual','Prioritario + sesión quincenal'],
['Ideal para','Inicio ordenado y una operación acotada.','Equipos de servicio con varias órdenes y responsables.','Empresa consolidada que necesita gobierno documental.']
],[1.55,1.75,1.75,1.75])
h2(doc,'Módulos incluidos en las tres opciones')
table(doc,['Bloque','Capacidades'],[
['Personas y habilitación','Personas fijas, por proyecto, disponibles y restringidas; ficha única, documentos, turnos y asistencia, EPP, formación, exámenes, aptitudes y salud ocupacional.'],
['Relación comercial','Clientes, contactos, contratos y firmas, órdenes de servicio, vigencias, responsables y bitácora.'],
['Cumplimiento','Documentación corporativa, habilitación del cliente, alertas, auditoría documental e historial de cambios.'],
['Gestión','Panel General, reportes base, configuración de empresa, importar/exportar, usuarios, permisos, privacidad y bitácora.']
],[2.0,5.05])
callout(doc,'Valor que se vende','“Una sola vista para saber quién está habilitado, qué documento vence y qué orden puede comenzar. Menos búsqueda, menos doble registro y más continuidad.”')
page(doc)

title(doc,'3. Nexo Klar Platinum','Operación conectada, recursos bajo control y decisiones con evidencia')
body(doc,'La línea Platinum incorpora la coordinación completa de recursos y ejecución. Conserva todos los módulos de Inicial y agrega activos, inventario, terceros, estadías, comunicaciones, incidentes, costos y reportes avanzados.')
table(doc,['Opción','Operación','Control','Corporativo'],[
['Precio mensual neto CLP','$649.000','$949.000','$1.490.000'],
['Personas activas incluidas','Hasta 100','Hasta 250','Hasta 500'],
['Usuarios nominados','12','25','50'],
['Órdenes de servicio activas/mes','100','300','Ilimitadas*'],
['Almacenamiento incluido','150 GB','500 GB','1 TB'],
['Soporte','Prioritario','Prioritario + revisión mensual','Gestor de éxito + revisión ejecutiva'],
['Ideal para','Operación en terreno con recursos y terceros.','Varias áreas, sedes o clientes activos.','Operación multiempresa o corporativa con integraciones.']
],[1.55,1.75,1.75,1.75])
body(doc,'*Uso razonable sujeto a política de servicio, almacenamiento y consumo contratado.')
h2(doc,'Módulos que se suman a Nexo Klar Inicial')
table(doc,['Bloque','Capacidades'],[
['Centro Operativo','Preparación de órdenes, brechas, responsables, alertas y coordinación diaria.'],
['Recursos y movilidad','Comunicaciones y convocatorias, vehículos, activos, equipos, credenciales, alojamientos y estadías.'],
['Terceros','Terceros y subcontratos, convenios, personal externo, habilitación, documentos y evaluación de desempeño.'],
['Activos e inventario','Maquinaria, instrumentos, herramientas, materiales, consumibles, bodegas, movimientos, mantenimiento, asignaciones y préstamos.'],
['Ejecución y control','Incidentes y no conformidades, acciones correctivas, reportes operacionales, costos y analítica avanzada según opción.']
],[2.0,5.05])
callout(doc,'Valor que se vende','“Antes de movilizar una orden, la empresa ve si tiene personas habilitadas, recursos disponibles, EPP, vehículo, alojamiento, documentación y evidencia para ejecutar.”')
page(doc)

title(doc,'4. Servicios adicionales independientes','Complementos de alto valor sin recargar el plan principal')
table(doc,['Servicio','Qué incluye','Precio mensual neto CLP','Condición'],[
['Libro de Obra Digital','Folios correlativos, anotaciones, acuerdos, instrucciones, evidencia, responsables, plazos, firmas y conexión con cliente, contrato y orden.','Desde $95.000','Firma electrónica, SMS y WhatsApp se cobran por consumo del proveedor.'],
['Gestor Comercial','Prospectos, contactos, bitácora, próxima acción, montos, documentos, pipeline y conversión a cliente, contrato u orden.','Desde $75.000','Incluye hasta 5 usuarios comerciales; usuarios extra según plan.'],
['Almacenamiento adicional','Bloques de capacidad para documentos, fotografías, evidencias y archivos históricos.','$25.000 / 50 GB','Se suman al plan activo.'],
['OCR / validación documental','Lectura de documentos, clasificación, extracción de fechas y revisión de excepciones.','Desde $95.000 + consumo','Se activa con proveedor y reglas aprobadas.'],
['Firma, mensajería e integración','Conectores a firma electrónica, correo, WhatsApp, ERP o API.','Desde $450.000 único + soporte','Cotización según proveedor, alcance, volumen y seguridad.']
],[1.4,2.9,1.2,1.55])
h2(doc,'Implementación y carga inicial')
table(doc,['Alternativa','Alcance','Precio neto CLP'],[
['Autoguiada','Plantillas, sesión de inicio de 90 minutos y revisión de primera carga.','Sin costo con anualidad; $190.000 en mensual.'],
['Asistida','Hasta 60 personas, 2 clientes, 3 contratos/órdenes, configuración base y dos sesiones de capacitación.','$590.000 único'],
['Operacional','Hasta 150 personas, 5 clientes, 10 contratos/órdenes, recursos, catálogo inicial, matriz y cuatro sesiones.','$1.290.000 único'],
['Corporativa / histórica','Diagnóstico, depuración, OCR, lotes, migración de cinco años, capacitación por equipo e informe de calidad.','Desde $3.500.000; se cotiza por volumen']
],[1.6,4.15,1.3])
callout(doc,'Regla de venta','Un piloto de 60 días puede cobrarse a $390.000 y acreditarse al 100% contra una contratación anual. El piloto siempre debe tener alcance, responsable, métrica y fecha de cierre.')
page(doc)

title(doc,'5. Árbol de decisión e integración','La explicación debe partir por las personas, no por las pantallas')
body(doc,'Esta es la ruta sugerida para presentar Nexo Klar. Cada nivel hereda el contexto del anterior y evita información duplicada. Se puede explicar en menos de cinco minutos durante una demo.')
table(doc,['Nivel','Pregunta del cliente','Módulos que responden','Resultado'],[
['1. Personas','¿Quién puede trabajar hoy?','Personas, tipo de trabajador, ficha, documentos, EPP, formación, exámenes, salud, turnos y restricciones.','Persona disponible, pendiente o restringida.'],
['2. Habilitación','¿Cumple lo necesario para esta orden?','Habilitación del cliente, auditoría, matriz de requisitos, alertas, credenciales y documentación.','Cumplimiento validado y vencimientos visibles.'],
['3. Asignación','¿Dónde y con quién trabaja?','Cliente, contrato, orden de servicio, asignación por proyecto, comunicaciones y convocatorias.','Persona vinculada a un servicio con responsable y fechas.'],
['4. Recursos','¿Tiene lo necesario para ejecutar?','EPP, vehículo, activo, herramienta, inventario, alojamiento, credencial y tercero.','Brechas operativas identificadas antes de movilizar.'],
['5. Ejecución','¿Qué ocurrió durante el trabajo?','Centro Operativo, asistencia, incidentes, acciones correctivas y Libro de Obra adicional.','Evidencia, compromisos, costos y estado actualizado.'],
['6. Decisión','¿Qué debemos corregir, renovar o vender?','Panel, alertas, reportes, analítica y Gestor Comercial adicional.','Decisiones por persona, cliente, contrato, orden o recurso.']
],[.75,1.6,3.05,1.6])
h2(doc,'Árbol del plan Inicial')
bullets(doc,[
'Personas → ficha, tipo de trabajador, turnos, documentos, EPP, formación, exámenes, salud y restricciones.',
'Habilitación → auditoría, alertas, credenciales y requisitos del cliente.',
'Asignación → cliente → contrato → orden de servicio → persona habilitada.',
'Control → Panel General, alertas y reportes base para actuar a tiempo.'
])
h2(doc,'Árbol del plan Platinum')
bullets(doc,[
'Todo Inicial → personas habilitadas y órdenes con contexto comercial.',
'Recursos → vehículo, activo, herramienta, bodega, EPP, alojamiento, credencial y tercero.',
'Ejecución → Centro Operativo → asistencia → incidentes / acciones correctivas → evidencia.',
'Decisión → reportes operacionales, costos, analítica, renovación y expansión de servicios.'
])
page(doc)

title(doc,'6. Matriz de visibilidad por línea','El cliente ve solo lo que necesita para operar')
table(doc,['Grupo funcional','Inicial','Platinum','Adicional'],[
['Panel y alertas','Panel General, alertas y reportes base.','Indicadores operacionales, costos y analítica avanzada.','—'],
['Personas y cumplimiento','Incluido.','Incluido + matrices y reglas avanzadas según opción.','OCR y validación documental.'],
['Relación comercial','Clientes, contratos, firmas y órdenes.','Incluido.','Gestor Comercial.'],
['Recursos y activos','EPP y control base de persona.','Vehículos, estadías, activos, inventario, mantenimiento y préstamos.','Almacenamiento adicional.'],
['Terceros','—','Subcontratos, convenios, personal, habilitaciones y evaluación.','—'],
['Ejecución','Seguimiento de orden, alertas y reportes base.','Centro Operativo, comunicaciones, incidentes y acciones correctivas.','Libro de Obra Digital.'],
['Gobierno','Usuarios, permisos, privacidad, bitácora e importar/exportar.','Configuración avanzada, costos e integración según opción.','API, ERP, firma y mensajería.']
],[2.0,1.65,2.2,1.2])
body(doc,'La Administración de clientes es de uso exclusivo del equipo interno de Nexo Klar. No forma parte de los planes vendidos a clientes finales.')
h2(doc,'Cómo recomendar una opción en la reunión')
table(doc,['Situación detectada','Recomendación'],[
['La empresa necesita ordenar personas, documentos, vencimientos, clientes y órdenes.','Nexo Klar Inicial Inicio o Equipo.'],
['Tiene varias órdenes, personal temporal, requisitos por cliente y más de un responsable.','Nexo Klar Inicial Empresa o Nexo Klar Platinum Operación.'],
['Coordina vehículos, hoteles, activos, subcontratos, EPP, bodegas e incidentes.','Nexo Klar Platinum Operación o Control.'],
['Tiene múltiples sedes, gran volumen, exigencias de auditoría, costos e integración.','Nexo Klar Platinum Corporativo.'],
['Necesita formalizar acuerdos de terreno o administrar el embudo antes del contrato.','Agregar Libro de Obra Digital y/o Gestor Comercial.']
],[3.1,3.95])
page(doc)

title(doc,'7. Condiciones comerciales y estrategia de cierre','Proteger el valor sin volver compleja la decisión')
table(doc,['Criterio','Propuesta'],[
['Plazo mínimo','12 meses para tarifas anuales. Modalidad mensual disponible con 15% adicional.'],
['Descuento anual','10% sobre suscripción o bonificación parcial de onboarding; no acumular ambos beneficios.'],
['Crecimiento','Personas, usuarios, órdenes, almacenamiento y módulos adicionales se activan por contrato sin migración ni pérdida de historial.'],
['Excesos','Se informa al administrador y se regulariza con bloque de capacidad durante el siguiente ciclo; no se bloquea información crítica.'],
['Soporte','Diferenciado por opción. Requerimientos especiales, desarrollo, integración y carga histórica se cotizan como orden de trabajo.'],
['Renovación','Revisión trimestral de uso, valor, capacidad, módulos y oportunidades de ajuste.']
],[2.0,5.05])
h2(doc,'Guion comercial recomendado')
bullets(doc,[
'Primero preguntar: “¿Cuánto tiempo tarda hoy en saber si una persona y una orden están listas para trabajar?”',
'Luego demostrar una única ruta: persona → requisitos → cliente → contrato → orden → recursos → alerta → reporte.',
'Recomendar una opción por dolor y capacidad, no por cantidad de módulos.',
'Cerrar con una implementación definida: qué se carga, quién valida, qué indicador se medirá y cuándo estará la primera operación activa.',
'Ofrecer Libro de Obra y Gestor Comercial únicamente cuando el problema lo justifique: formalidad de ejecución o control preventa.'
])
callout(doc,'Mensaje de cierre','“Nexo Klar no obliga a la empresa a cambiar todo de una vez. Parte con el control que necesita hoy y crece hacia la operación conectada cuando el negocio lo requiere.”')

h2(doc,'Referencias de modelo comercial')
body(doc,'Esta estructura considera patrones vigentes de plataformas internacionales de field service y gestión operacional: Housecall Pro combina planes por usuarios y acompañamiento; ServiceM8 combina capacidad mensual de trabajos, mensajería y complementos; FieldPulse publica una lógica por técnico. Los valores de Nexo Klar son una propuesta propia para Chile y deben validarse con pilotos, costos cloud, soporte y consumo real.')
bullets(doc,[
'Housecall Pro Pricing: niveles por usuarios, soporte y plan anual.',
'ServiceM8 Pricing: capacidad de trabajos, mensajería, complementos y crecimiento de plan.',
'FieldPulse Pricing: modelo por técnico y paquetes configurables.'
])
callout(doc,'Resumen','Nexo Klar Inicial ordena y habilita. Nexo Klar Platinum coordina, controla y escala. Libro de Obra Digital y Gestor Comercial amplían la trazabilidad sin obligar a todos los clientes a pagar por capacidades que no usan.')

OUT.parent.mkdir(parents=True,exist_ok=True); REPO_COPY.parent.mkdir(parents=True,exist_ok=True)
doc.save(OUT); doc.save(REPO_COPY)
print(OUT); print(REPO_COPY)
