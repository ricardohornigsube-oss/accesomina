from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

IMAGE = Path('/Users/ricardo.hornig/Downloads/WhatsApp Image 2026-08-17 at 16.30.13.jpeg')
OUT = Path('/Users/ricardo.hornig/Desktop/Nexo Klar/Nexo_Klar_Logo_Pulso_Referencia.docx')
REPO = Path('/Users/ricardo.hornig/Desktop/Acceso Mina /Acceso Mina/docs/Nexo_Klar_Logo_Pulso_Referencia.docx')

doc = Document()
section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width, section.page_height = section.page_height, section.page_width
section.top_margin = Inches(.48); section.bottom_margin = Inches(.48)
section.left_margin = Inches(.48); section.right_margin = Inches(.48)

styles = doc.styles
styles['Normal'].font.name = 'Arial'; styles['Normal'].font.size = Pt(10)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('NEXO KLAR · LOGO PULSO'); r.bold = True; r.font.name = 'Arial'; r.font.size = Pt(17); r.font.color.rgb = RGBColor(42,42,140)
p.paragraph_format.space_after = Pt(3)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Concepto, logo horizontal, isotipo, logo vertical y aplicaciones'); r.font.name='Arial'; r.font.size=Pt(9); r.font.color.rgb=RGBColor(93,107,122)
p.paragraph_format.space_after = Pt(10)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run().add_picture(str(IMAGE), width=Inches(10.0))
p.paragraph_format.space_after = Pt(8)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Referencia visual recibida por WhatsApp · El concepto “Flujo inteligente” se conserva como base de identidad.'); r.font.name='Arial'; r.font.size=Pt(8); r.font.color.rgb=RGBColor(93,107,122)

OUT.parent.mkdir(parents=True,exist_ok=True); REPO.parent.mkdir(parents=True,exist_ok=True)
doc.save(OUT); doc.save(REPO)
print(OUT); print(REPO)
