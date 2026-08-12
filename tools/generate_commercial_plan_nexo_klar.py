#!/usr/bin/env python3
"""Create Nexo Klar commercial plan in Word with visual verification support."""
from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / 'docs' / 'Plan_Comercial_Nexo_Klar_2026_2027.docx'
IMG = ROOT / 'docs' / '_plan_comercial_nexo_klar_flujo.png'

NAVY='141A20'; INDIGO='2A2A8C'; PRIMARY=INDIGO; DEEP='1A1A5E'; TEAL='00CFC1'; TEAL_D='00706A'; MUTED='5D6B7A'; LINE='E3DED2'; PALE='F0F0FA'; ALT='FBF9F5'; GREEN='1B7F4B'; AMBER='C77700'; RED='B3261E'

def hexrgb(value): return RGBColor.from_string(value)
def shade(cell, color):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),color); tcPr.append(shd)
def border(cell, color=LINE, size='8'):
    tcPr=cell._tc.get_or_add_tcPr(); borders=tcPr.first_child_found_in('w:tcBorders')
    if borders is None: borders=OxmlElement('w:tcBorders');tcPr.append(borders)
    for edge in ('top','left','bottom','right'):
        tag='w:'+edge;el=borders.find(qn(tag))
        if el is None: el=OxmlElement(tag);borders.append(el)
        el.set(qn('w:val'),'single');el.set(qn('w:sz'),size);el.set(qn('w:color'),color)
def prevent_row_split(row):
    trPr=row._tr.get_or_add_trPr();el=OxmlElement('w:cantSplit');trPr.append(el)
def set_cell(cell, text, bold=False, color=NAVY, size=9, align=None):
    cell.text='';p=cell.paragraphs[0]
    if align is not None:p.alignment=align
    r=p.add_run(text);r.bold=bold;r.font.name='Arial';r.font.size=Pt(size);r.font.color.rgb=hexrgb(color)
    cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for cell_p in cell.paragraphs:
        cell_p.paragraph_format.space_after=Pt(3);cell_p.paragraph_format.space_before=Pt(3)
def add_title(doc, text, subtitle=None):
    p=doc.add_paragraph();p.paragraph_format.space_after=Pt(6)
    r=p.add_run(text);r.bold=True;r.font.name='Arial';r.font.size=Pt(24);r.font.color.rgb=hexrgb(NAVY)
    if subtitle:
        p=doc.add_paragraph();p.paragraph_format.space_after=Pt(12)
        r=p.add_run(subtitle);r.font.name='Arial';r.font.size=Pt(11);r.font.color.rgb=hexrgb(MUTED)
def add_heading(doc, text):
    p=doc.add_paragraph();p.paragraph_format.space_before=Pt(12);p.paragraph_format.space_after=Pt(6)
    r=p.add_run(text);r.bold=True;r.font.name='Arial';r.font.size=Pt(15);r.font.color.rgb=hexrgb(DEEP)
    return p
def add_body(doc,text,bold_prefix=None):
    p=doc.add_paragraph();p.paragraph_format.space_after=Pt(6);p.paragraph_format.line_spacing=1.15
    if bold_prefix and text.startswith(bold_prefix):
        r=p.add_run(bold_prefix);r.bold=True;r.font.name='Arial';r.font.size=Pt(10);r.font.color.rgb=hexrgb(NAVY);text=text[len(bold_prefix):]
    r=p.add_run(text);r.font.name='Arial';r.font.size=Pt(10);r.font.color.rgb=hexrgb(NAVY)
    return p
def add_bullets(doc,items):
    for item in items:
        p=doc.add_paragraph(style='List Bullet');p.paragraph_format.space_after=Pt(3);p.paragraph_format.left_indent=Inches(.25)
        r=p.add_run(item);r.font.name='Arial';r.font.size=Pt(9.5);r.font.color.rgb=hexrgb(NAVY)
def add_callout(doc,title,text,color=TEAL_D):
    t=doc.add_table(rows=1,cols=2);t.alignment=WD_TABLE_ALIGNMENT.CENTER;t.autofit=False
    t.columns[0].width=Inches(1.35);t.columns[1].width=Inches(5.75)
    c1,c2=t.rows[0].cells;shade(c1,'EAFBF9');shade(c2,'EAFBF9');border(c1,'BDEDE8');border(c2,'BDEDE8')
    set_cell(c1,title,True,color,9);set_cell(c2,text,False,NAVY,9)
    doc.add_paragraph().paragraph_format.space_after=Pt(1)
def add_table(doc,headers,rows,widths=None):
    t=doc.add_table(rows=1,cols=len(headers));t.alignment=WD_TABLE_ALIGNMENT.CENTER;t.autofit=False
    for i,h in enumerate(headers):
        if widths:t.columns[i].width=Inches(widths[i])
        shade(t.rows[0].cells[i],DEEP);border(t.rows[0].cells[i],DEEP);set_cell(t.rows[0].cells[i],h,True,'FFFFFF',8.5)
    for idx,row in enumerate(rows):
        cells=t.add_row().cells
        for i,v in enumerate(row):
            if widths:cells[i].width=Inches(widths[i])
            shade(cells[i],ALT if idx%2==0 else 'FFFFFF');border(cells[i]);set_cell(cells[i],v,False,NAVY,8.5)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
    return t
def add_plan_card(doc, plan, focus, price, annual, scale, modules, value):
    t=doc.add_table(rows=1,cols=2);t.alignment=WD_TABLE_ALIGNMENT.CENTER;t.autofit=False
    t.columns[0].width=Inches(1.55);t.columns[1].width=Inches(5.55)
    left,right=t.rows[0].cells;header=left.merge(right);shade(header,DEEP);border(header,DEEP)
    set_cell(header,plan+'  |  '+price,True,'FFFFFF',12)
    prevent_row_split(t.rows[0])
    for label,content in [('Enfoque',focus),('Escala incluida',scale),('Módulos principales',modules),('Valor que habilita',value)]:
        cells=t.add_row().cells
        cells[0].width=Inches(1.55);cells[1].width=Inches(5.55)
        shade(cells[0],PALE);shade(cells[1],'FFFFFF');border(cells[0]);border(cells[1])
        set_cell(cells[0],label,True,DEEP,8.5)
        set_cell(cells[1],content,False,NAVY,8.8)
        prevent_row_split(t.rows[-1])
    p=doc.add_paragraph();p.paragraph_format.space_after=Pt(5)
    r=p.add_run('Pago anual anticipado: '+annual+' netos.');r.bold=True;r.font.name='Arial';r.font.size=Pt(8.5);r.font.color.rgb=hexrgb(TEAL_D)
    return t
def page_break(doc): doc.add_page_break()
def add_centered_picture(doc, path, width):
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path),width=width)
    return p
def header_footer(section):
    h=section.header.paragraphs[0];h.alignment=WD_ALIGN_PARAGRAPH.RIGHT;r=h.add_run('NEXO KLAR SPA · PLAN COMERCIAL');r.font.name='Arial';r.font.size=Pt(7.5);r.font.color.rgb=hexrgb(MUTED)
    f=section.footer.paragraphs[0];f.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=f.add_run('Confidencial · Plan comercial 2026–2027');r.font.name='Arial';r.font.size=Pt(7.5);r.font.color.rgb=hexrgb(MUTED)

def make_diagram():
    w,h=1600,650; im=Image.new('RGB',(w,h),'#FBF9F5');d=ImageDraw.Draw(im)
    try: title=ImageFont.truetype('/System/Library/Fonts/Supplemental/Arial Bold.ttf',34); small=ImageFont.truetype('/System/Library/Fonts/Supplemental/Arial.ttf',20); boxfont=ImageFont.truetype('/System/Library/Fonts/Supplemental/Arial Bold.ttf',22)
    except: title=small=boxfont=None
    d.text((70,45),'Nexo Klar: flujo comercial y operacional integrado',fill='#141A20',font=title)
    boxes=[('Prospecto','Oportunidad y\nseguimiento'),('Cliente','Datos, contactos\ny contexto'),('Contrato','Vigencia, alcance\ny firmas'),('Orden de servicio','Planificación,\npersonas y recursos'),('Ejecución','Evidencia, libro,\nincidentes y costos'),('Control','Alertas, auditoría,\nreportes y mejora')]
    x=60;y=210;bw=220;bh=150;gap=35
    for i,(a,b) in enumerate(boxes):
        xx=x+i*(bw+gap); d.rounded_rectangle((xx,y,xx+bw,y+bh),radius=18,fill='#FFFFFF',outline='#2A2A8C',width=4)
        d.text((xx+20,y+28),a,fill='#141A20',font=boxfont)
        d.multiline_text((xx+20,y+70),b,fill='#5D6B7A',font=small,spacing=4)
        if i<len(boxes)-1:
            d.line((xx+bw,y+75,xx+bw+gap-8,y+75),fill='#00AFA5',width=5);d.polygon([(xx+bw+gap-8,y+75),(xx+bw+gap-24,y+66),(xx+bw+gap-24,y+84)],fill='#00AFA5')
    d.rounded_rectangle((120,455,1480,570),radius=16,fill='#EEF6FF',outline='#AAB7F0',width=2)
    d.text((155,482),'Capas que acompañan todo el flujo:',fill='#2A2A8C',font=boxfont)
    d.text((155,520),'Personas · cumplimiento documental · EPP · turnos · activos e inventario · terceros · comunicaciones · privacidad · usuarios y permisos',fill='#5D6B7A',font=small)
    im.save(IMG)

make_diagram()
doc=Document();sec=doc.sections[0];sec.top_margin=Inches(.62);sec.bottom_margin=Inches(.62);sec.left_margin=Inches(.72);sec.right_margin=Inches(.72);header_footer(sec)

# Cover
doc.add_paragraph().paragraph_format.space_after=Pt(45)
p=doc.add_paragraph();r=p.add_run('NEXO KLAR SPA');r.bold=True;r.font.name='Arial';r.font.size=Pt(11);r.font.color.rgb=hexrgb(TEAL_D)
p=doc.add_paragraph();p.paragraph_format.space_after=Pt(10);r=p.add_run('Plan Comercial\n2026–2027');r.bold=True;r.font.name='Arial';r.font.size=Pt(31);r.font.color.rgb=hexrgb(NAVY)
p=doc.add_paragraph();p.paragraph_format.space_after=Pt(28);r=p.add_run('Funcionalidad, propuesta de valor, ventajas competitivas y estrategia para vender una plataforma SaaS de control operacional.');r.font.name='Arial';r.font.size=Pt(13);r.font.color.rgb=hexrgb(MUTED)
add_callout(doc,'Objetivo','Transformar la funcionalidad existente en una oferta comercial simple: demostrar control, reducir riesgo operativo, acelerar la preparación de servicios y conservar la información dentro de cada empresa.',TEAL_D)
add_centered_picture(doc,IMG,Inches(6.8))
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=p.add_run('Documento interno para dirección, ventas y pilotos comerciales · 12 de agosto de 2026');r.font.name='Arial';r.font.size=Pt(8.5);r.font.color.rgb=hexrgb(MUTED)
page_break(doc)

add_title(doc,'1. Resumen ejecutivo','Qué se vende y por qué importa')
add_body(doc,'Nexo Klar es una plataforma SaaS multiempresa para ordenar y conectar la operación de empresas de servicios. Reúne la relación comercial, las personas, los recursos, el cumplimiento, la ejecución y los reportes en un solo espacio privado por empresa.')
add_table(doc,['Aspecto','Síntesis comercial'],[
 ['Problema del cliente','Información repartida entre planillas, correos, carpetas, teléfonos y personas; vencimientos y recursos se controlan tarde.'],
 ['Propuesta de valor','Una fuente de información operativa conectada desde el prospecto hasta el cierre del servicio.'],
 ['Diferenciación','No se limita a documentos: integra personas, EPP, recursos, inventario, hotelería, turnos, terceros, incidentes, libro de obra y analítica.'],
 ['Resultado esperado','Menos horas administrativas, menor riesgo por vencimientos, mayor trazabilidad y continuidad cuando cambian los equipos.'],
 ['Modelo comercial','Suscripción SaaS por empresa + implementación + adicionales de alto valor.']
],[1.55,5.25])
add_callout(doc,'Mensaje central','“Nexo Klar ayuda a que la operación no dependa de planillas ni de una persona. Cada dato queda conectado, vigente y disponible para quien debe decidir o ejecutar.”',PRIMARY)

add_title(doc,'2. Revisión de funcionalidad y usabilidad','Valor real que hoy puede demostrarse')
add_body(doc,'La plataforma cuenta con una base funcional amplia y conectada. La revisión automática actual registra 123 pruebas correctas, incluyendo relaciones entre módulos, aislamiento multiempresa, validación de duplicados, archivos, inventario, libro de obra, importación y exportación.')
add_table(doc,['Bloque funcional','Qué resuelve para el cliente','Beneficio visible'],[
 ['Relación comercial','Prospectos, clientes, contratos, firmas y órdenes de servicio.','Un contexto comercial único y trazable.'],
 ['Personas y cumplimiento','Trabajadores fijos, por proyecto y disponibles; documentos, cursos, exámenes, salud, credenciales y EPP.','Saber quién está habilitado, pendiente o restringido.'],
 ['Operación y recursos','Turnos, asistencia, vehículos, hotelería, comunicaciones, terceros, activos, bodegas e inventario.','Planificar y movilizar personas y recursos con menos improvisación.'],
 ['Ejecución y calidad','Centro operativo, libro de obra, incidentes, no conformidades, CAPA y auditoría.','Registrar evidencia, compromisos y acciones correctivas.'],
 ['Gobierno y decisión','Alertas, reportes, permisos, bitácora, privacidad, configuración e importación/exportación.','Control ejecutivo, continuidad y capacidad de auditoría.']
],[1.55,3.5,1.75])
add_body(doc,'<b>Lectura de usabilidad.</b> La principal fortaleza es que el usuario puede iniciar desde un cliente, contrato u orden de servicio y llegar a las personas, documentos, recursos y evidencia relacionados. Esto reduce búsquedas y evita mantener registros paralelos.',bold_prefix='<b>Lectura de usabilidad.</b>')
page_break(doc)

add_title(doc,'3. Ventajas y beneficios comerciales','Por qué un cliente compraría Nexo Klar')
add_table(doc,['Ventaja','Beneficio para el cliente','Prueba en demo'],[
 ['Información conectada','Evita duplicar datos entre RR.HH., operaciones, prevención y administración.','Abrir una orden y mostrar personas, documentos, hotel, vehículo, EPP y bitácora.'],
 ['Control de cumplimiento','Detecta vencimientos, requisitos faltantes, observaciones y restricciones.','Panel de alertas agrupadas por persona, recurso o servicio.'],
 ['Gestión de servicios','Une clientes, contratos, órdenes, dotación, recursos y costos.','Crear orden de servicio y asignar recursos disponibles.'],
 ['Trazabilidad','Conserva cambios, evidencias, firmas, compromisos y resultados.','Mostrar bitácora, auditoría, libro de obra e historial documental.'],
 ['Multiempresa segura','Cada cliente tiene usuarios, marca, catálogos y datos independientes.','Cambiar entre entornos de empresa y comprobar que no se mezclan datos.'],
 ['Adaptable a industrias','La nomenclatura es neutral y los catálogos se parametrizan por empresa.','Demo por industria: servicios técnicos, construcción, energía, logística o minería.'],
 ['Escalabilidad comercial','Permite habilitar módulos y adicionales según el plan contratado.','Mostrar plan Esencial, Operación e Integral sin cambiar la base de datos.']
],[1.35,3.3,2.15])
add_callout(doc,'Ventaja competitiva','Competidores locales suelen enfocarse en acreditación/documentación de contratistas o en gestión de terreno. Nexo Klar puede posicionarse en la intersección: control documental + ejecución de servicios + recursos físicos + trazabilidad comercial.',TEAL_D)

page_break(doc)
add_title(doc,'4. Mercado y posicionamiento','Dónde competir primero')
add_body(doc,'La revisión de mercado muestra alternativas chilenas enfocadas en cumplimiento documental y contratistas, como HabiPass, Verify, SubCheck, PreveSafe y Fácil Control; y herramientas de field service como Praxedo. Esta evidencia confirma una oportunidad de posicionar Nexo Klar como plataforma integrada para empresas de servicios, sin depender únicamente de un sector.')
add_table(doc,['Segmento','Dolor dominante','Posicionamiento Nexo Klar'],[
 ['Empresas de mantención industrial','Planificación de cuadrillas, acreditación, herramientas, costos y evidencia de servicio.','Control operacional de la orden de servicio.'],
 ['Construcción y montaje','Subcontratos, personas, turnos, documentación, avances e incidentes.','Trazabilidad de obra, recursos y cumplimiento.'],
 ['Energía y utilities','Permisos, cuadrillas, equipos, seguridad y contratos de servicios.','Coordinación de servicios críticos y habilitación.'],
 ['Facilities y multiservicios','Personal externo, SLA, credenciales, activos y reportes por cliente.','Operación repetitiva por instalación y cliente.'],
 ['Logística y servicios técnicos','Turnos, vehículos, equipos, personal, documentos y respaldo de ejecución.','Despacho, recursos, evidencias y continuidad.'],
 ['Minería','Acreditación, exigencias de mandante, faena, salud, EPP y subcontratos.','Caso de uso profundo, no único mercado.']
],[1.65,2.3,2.85])
page_break(doc)

add_title(doc,'5. Cliente ideal y decisores','A quién vender primero')
add_table(doc,['Perfil','Características','Decisor / usuario'],[
 ['Cliente ideal inicial','Empresa de servicios con 30 a 300 personas, varias órdenes activas, personal temporal y presión por documentos o recursos.','Gerente de Operaciones, Administrador de Contrato, Gerencia General.'],
 ['Cliente con dolor urgente','Trabaja con varios clientes/mandantes, recibe observaciones frecuentes, usa planillas y necesita demostrar evidencia.','Prevención, Acreditación, RR.HH., Calidad.'],
 ['Cliente de expansión','Opera múltiples zonas, subcontratos, bodegas, vehículos o alojamiento.','Jefatura de Operaciones, Finanzas, Logística.'],
 ['Influyente interno','Necesita usar el sistema día a día y validar su valor.','Coordinador de servicios, reclutador, encargado documental, prevencionista.']
],[1.6,3.3,1.9])
add_body(doc,'La venta debe comenzar con un problema verificable y no con una lista de módulos: “¿Cuánto tiempo tardan en saber si una persona, vehículo o servicio está habilitado para ejecutar?”. Luego se demuestra el flujo conectado y se mide un indicador inicial.')
add_callout(doc,'Oferta de entrada','Vender un piloto de 60 a 90 días con una operación acotada, una persona responsable del cliente, carga inicial guiada y objetivos medibles de tiempo, cumplimiento y trazabilidad.',PRIMARY)

add_title(doc,'6. Referencia de mercado y criterio de precio','Precio transparente, valor protegido y negociación simple')
add_body(doc,'En Chile, gran parte de las plataformas de acreditación y gestión de contratistas cotiza según alcance, personas, sedes, documentos e implementación. En field service internacional se observa una combinación de suscripción por usuario, niveles funcionales, complementos y acompañamiento pagado. Nexo Klar debe ocupar esa lógica: una tarifa base clara, crecimiento por uso y servicios de implementación visibles.')
add_table(doc,['Referencia pública','Modelo comercial observado','Implicancia para Nexo Klar'],[
 ['FieldPulse','Precio por usuario y cotización personalizada por paquete; complementos de IA, flota, formularios y soporte.','Cobrar por escala y no esconder el valor de módulos o soporte especializado.'],
 ['Housecall Pro','Planes públicos desde USD 59/mes y niveles superiores; suma usuarios y onboarding especializado en planes altos.','Mantener un plan de entrada simple y usar soporte/acompañamiento como diferenciador de niveles.'],
 ['ServiceM8','Inicio autoguiado o implementación y capacitación con socios especializados.','Ofrecer carga autoguiada y carga asistida como alternativas, no una sola tarifa rígida.'],
 ['Mercado chileno de contratistas','Plataformas de acreditación, cumplimiento y subcontratación normalmente solicitan cotización según empresa y alcance.','Posicionar Nexo Klar por valor integrado y cotización empresarial, manteniendo precios referenciales públicos.']
],[1.45,2.65,3.2])
add_callout(doc,'Criterio comercial','No competir solo por precio. El precio base debe abrir la conversación; la implementación, los activos adicionales, las integraciones y el soporte reflejan el esfuerzo real y se cotizan por alcance.',PRIMARY)

page_break(doc)
add_title(doc,'7. Oferta SaaS recomendada','Tres planes claros, escalables y negociables')
add_body(doc,'Los precios se definieron con una lógica de valor: Nexo Klar no reemplaza una pantalla aislada, sino que conecta relación comercial, personas, cumplimiento, recursos, ejecución y control. La tarifa se ajusta por dotación activa y complejidad operacional, no por cantidad de pantallas abiertas.')
add_plan_card(doc,'Nexo Klar Base','Ordenar la operación y dejar de depender de planillas.','$249.000 / mes','$2.689.200','Hasta 30 personas activas y 5 usuarios nominados.','Panel, clientes, contratos, órdenes de servicio, personas, alertas, documentos, reportes base, usuarios, configuración e importación/exportación.','Una fuente única para conocer el estado de clientes, contratos, órdenes, personas y vencimientos.')
add_plan_card(doc,'Nexo Klar Operación','Coordinar personas y recursos en una operación de servicios recurrente.','$490.000 / mes','$5.292.000','Hasta 75 personas activas y 10 usuarios nominados.','Todo Base + gestión de personal por proyecto, turnos, EPP, formación, exámenes, comunicaciones, vehículos, estadías, terceros, incidentes y Centro Operativo.','Preparar servicios con mayor anticipación y reducir brechas de personas, documentos, recursos y coordinación.')
add_callout(doc,'Valor por usuario','A capacidad completa, el costo mensual equivale a $49.800 por usuario Base, $49.000 por usuario Operación y $44.500 por usuario Integral. Este indicador sirve para explicar el valor de una herramienta compartida por equipos administrativos y operativos.',TEAL_D)
page_break(doc)
add_plan_card(doc,'Nexo Klar Integral','Gobierno operacional, cumplimiento avanzado y crecimiento multioperación.','$890.000 / mes','$9.612.000','Hasta 200 personas activas y 20 usuarios nominados.','Todo Operación + auditoría avanzada, habilitación del cliente, activos e inventario, costos, rentabilidad, automatizaciones, portal, API e identidad visual por empresa.','Entregar trazabilidad para gestión, auditoría, costos, expansión y decisiones ejecutivas.')
add_table(doc,['Crecimiento y adicionales','Precio lista neto CLP','Regla comercial'],[
 ['Persona activa adicional','Base: $2.500 · Operación: $3.500 · Integral: $5.000 / mes','Cobrar solo sobre la dotación activa promedio del mes.'],
 ['Usuario adicional','Base: $18.000 · Operación: $25.000 · Integral: $35.000 / mes','Solo para usuarios con acceso nominativo que excedan el plan.'],
 ['Libro de obra digital','Desde $95.000 / mes','Adicional por empresa; firma electrónica y mensajes se cobran según consumo del proveedor.'],
 ['Prospectos y oportunidades','Desde $65.000 / mes','Adicional comercial para conectar preventa, seguimiento y conversión.'],
 ['Integración, API o desarrollo específico','Desde $450.000 único + soporte mensual según alcance','Se cotiza después de una sesión de descubrimiento y no se promete como estándar.'],
 ['Almacenamiento extraordinario, OCR, firma o mensajería','Costo de proveedor + administración Nexo Klar','Se informa por separado antes de activar el servicio.']
],[2.15,1.55,3.6])
add_body(doc,'Condición sugerida: contrato mínimo de 12 meses. Pago anual anticipado: 10% de descuento en la suscripción o bonificación parcial del onboarding, no ambos. Para clientes piloto, aplicar un descuento temporal y dejarlo expresamente limitado a los primeros 90 días. Banda de negociación recomendada: máximo 15% de descuento sobre tarifa lista, siempre a cambio de plazo anual, caso de éxito autorizado o volumen comprometido.')

page_break(doc)
add_title(doc,'8. Carga inicial y onboarding','Alternativas claras para no regalar trabajo crítico')
add_table(doc,['Alternativa','Alcance','Precio neto CLP','Cuándo ofrecerla'],[
 ['Autoguiada','Plantillas, centro de ayuda, sesión de inicio de 90 minutos y revisión de carga del cliente.','Sin costo en contrato anual; $150.000 en mensual.','Empresa pequeña, datos ordenados y administrador disponible.'],
 ['Asistida','Configuración base, carga de hasta 50 personas, 2 clientes, 3 contratos/órdenes, importación CSV y 2 sesiones de capacitación.','$490.000 único','Oferta estándar para la mayoría de empresas de servicios.'],
 ['Operacional','Catálogos, matriz inicial, carga de hasta 150 personas, 5 clientes, 10 contratos/órdenes, recursos y 4 sesiones de capacitación.','$990.000 único','Cliente con varias operaciones, personal temporal o necesidad de partir rápido.'],
 ['Corporativa','Diagnóstico, depuración de datos, migración mayor, roles, identidad visual, capacitación por equipo y plan de adopción.','Desde $2.500.000 + alcance','Empresa multioperación, alto volumen o requerimientos particulares.']
],[1.3,3.75,1.15,1.1])
add_bullets(doc,[
 'Alternativa de cierre: descontar hasta el 50% del onboarding asistido u operacional contra el primer año pagado por anticipado. Evita que la carga inicial parezca una barrera, sin eliminar su valor.',
 'Piloto recomendado: 60 días pagados a $350.000, acreditables al 100% a la implementación si se contrata un plan anual. No ofrecer pilotos ilimitados ni cargas masivas gratis.',
 'La carga histórica, corrección de datos, parametrizaciones especiales e integraciones deben quedar fuera del precio base y documentarse en una orden de trabajo.'
])
add_callout(doc,'Forma de vender la carga inicial','“Ustedes pueden cargar con plantillas y guía, o podemos dejar una operación real lista para trabajar. La diferencia es el alcance, el tiempo y el acompañamiento; la plataforma sigue siendo la misma.”',TEAL_D)

page_break(doc)
add_title(doc,'9. Estrategia de ventas','Proceso comercial repetible')
add_table(doc,['Etapa','Acción','Entregable / métrica'],[
 ['1. Generar demanda','Contenido en LinkedIn, videos cortos por industria, red de contactos, alianzas con prevención/RR.HH. y referidos.','Lista de cuentas objetivo y reuniones calificadas.'],
 ['2. Descubrimiento','Entrevista de 30–45 minutos: proceso actual, planillas, servicios, documentos, personas, recursos y riesgo.','Mapa de dolor y caso de uso prioritario.'],
 ['3. Demo contextual','Demostrar un solo flujo completo: cliente → contrato → orden → persona → recurso → evidencia → alerta → reporte.','Demo de 45 minutos y resumen de valor.'],
 ['4. Piloto','Configurar una empresa, cliente u orden real; cargar datos iniciales y medir indicadores.','Informe de resultados a 60–90 días.'],
 ['5. Propuesta','Plan recomendado, módulos, implementación, soporte, condiciones y próximos hitos.','Propuesta comercial simple, alcance y ROI cualitativo.'],
 ['6. Cierre y expansión','Onboarding, adopción, revisión mensual, activación de módulos y caso de éxito.','Renovación, upsell y recomendación.']
],[1.05,3.1,3.15])
add_callout(doc,'Regla de demo','No mostrar 30 pantallas. Mostrar un caso real con principio y fin: preparar un servicio, habilitar recursos, ejecutar, registrar evidencia y obtener un reporte.',TEAL_D)

add_title(doc,'10. Plan de marketing de 90 días','Construir confianza antes de escalar')
add_table(doc,['Periodo','Acciones principales','Resultado esperado'],[
 ['Días 1–30','Definir 3 industrias foco, página por caso de uso, demo con datos realistas, presentación comercial, guion de diagnóstico y 20 cuentas objetivo.','Mensaje claro y base de prospección.'],
 ['Días 31–60','Publicar 2 contenidos semanales, enviar mensajes personalizados, realizar demos, seleccionar 1–2 pilotos y levantar línea base.','Primeros pilotos y aprendizaje comercial.'],
 ['Días 61–90','Medir resultados del piloto, producir caso de éxito, preparar webinar o video por industria y ofertas de cierre anual.','Referencias, propuesta validada y primeros contratos SaaS.']
],[1.2,3.5,2.6])
add_bullets(doc,[
 'Contenido recomendado: “cómo saber si una orden está lista para ejecutar”, “cómo evitar que el conocimiento se vaya con una persona”, “cómo controlar documentos y recursos en una sola vista”.',
 'Activos comerciales mínimos: sitio público, video de 90 segundos, demo por industria, presentación de 10 diapositivas, propuesta estándar, piloto y caso de éxito.',
 'Canales iniciales: LinkedIn personal y corporativo, red de contratistas, asociaciones sectoriales, aliados de prevención y consultores de gestión.'
])
page_break(doc)

add_title(doc,'11. Métricas de venta y adopción','Cómo medir el plan')
add_table(doc,['Nivel','Indicadores propuestos'],[
 ['Demanda','Cuentas objetivo contactadas, tasa de respuesta, reuniones de descubrimiento, demos realizadas y costo por reunión.'],
 ['Conversión','Demo → piloto, piloto → contrato anual, ciclo comercial, ticket mensual, setup vendido y adicionales activados.'],
 ['Adopción','Usuarios activos, registros creados, órdenes gestionadas, documentos revisados, alertas resueltas y reportes emitidos.'],
 ['Valor operativo','Tiempo de preparación del servicio, vencimientos detectados antes de vencer, porcentaje de personas habilitadas, reducción de planillas y trazabilidad documental.'],
 ['Retención','Renovación, expansión de módulos, satisfacción del administrador, tickets resueltos y recomendación.']
],[1.45,5.85])
add_heading(doc,'Proyección orientativa de 12 meses')
add_table(doc,['Escenario','Clientes al mes 12','MRR estimado','Supuestos'],[
 ['Conservador','6–8','$1,5M–$2,2M CLP','Ventas directas, 1 piloto exitoso y mezcla Esencial/Operación.'],
 ['Objetivo','10–14','$3,0M–$4,8M CLP','Proceso comercial semanal, 2 casos de éxito y venta de módulos adicionales.'],
 ['Acelerado','18–25','$6,0M–$9,0M CLP','Alianzas, referidos, marketing por industria y foco en Operación/Integral.']
],[1.3,1.6,1.6,2.8])
add_body(doc,'Estas proyecciones son escenarios de planificación, no promesas de facturación. Deben revisarse mensualmente con tasas reales de reunión, demo, piloto, cierre y churn.')

add_title(doc,'12. Riesgos comerciales y cómo manejarlos','Vender con claridad')
add_table(doc,['Riesgo','Cómo tratarlo comercialmente'],[
 ['Prometer IA o firma como si ya estuvieran activas','Presentar OCR, firma, WhatsApp y validación documental como conectores preparados que requieren proveedor y configuración.'],
 ['Demasiados módulos en la primera demo','Partir con un proceso crítico y expandir por plan/adicionales.'],
 ['Implementación sin dueño del cliente','Exigir administrador responsable, alcance piloto y calendario de carga inicial.'],
 ['Precio desconectado del valor','Cobrar setup por carga/configuración y tarifa mensual según personas, módulos, soporte e integraciones.'],
 ['Dependencia de un cliente piloto','Usar el aprendizaje para construir plantillas de industria, no personalizaciones irrepetibles.'],
 ['Expectativa de producción inmediata','Usar checklist de habilitación cloud, seguridad e integraciones antes de comprometer SLA corporativo.']
],[2.3,4.7])
add_callout(doc,'Decisión recomendada','Iniciar ventas con un piloto acotado en mantenimiento industrial, construcción o facilities; conservar minería como caso de uso profundo. Consolidar 2 casos de éxito antes de comprometer una expansión masiva o integraciones complejas.',AMBER)

add_title(doc,'13. Próximos pasos de dirección','Plan de ejecución')
add_bullets(doc,[
 'Aprobar la oferta comercial de tres planes y los adicionales Libro de obra digital y Prospectos y oportunidades.',
 'Elegir tres industrias foco y preparar una demo específica para cada una usando la misma base funcional.',
 'Definir un piloto gratuito o con descuento, con alcance, indicadores y autorización para producir un caso de éxito.',
 'Cerrar el checklist de producción cloud: AWS, RDS, S3, antivirus, MFA, monitoreo, respaldos, firma, correo y WhatsApp.',
 'Asignar responsables: ventas, implementación, soporte, seguridad/privacidad y operación técnica.',
 'Revisar mensualmente pipeline, adopción, costos de soporte, aprendizaje de pilotos y módulos con mayor demanda.'
])
add_callout(doc,'Cierre','Nexo Klar ya tiene la amplitud funcional para vender control operacional integrado. El foco comercial debe ser hacer tangible el valor: menos búsqueda, menos riesgo, más continuidad y decisiones con información conectada.',PRIMARY)
add_heading(doc,'Referencias de mercado utilizadas')
add_bullets(doc,[
 'FieldPulse, página de precios: modelo por usuario, cotización por paquete y complementos operacionales (fieldpulse.com/pricing).',
 'Housecall Pro, página de precios: plan de entrada publicado desde USD 59/mes y acompañamiento especializado en niveles altos (housecallpro.com/pricing).',
 'ServiceM8, página de precios: alternativas de inicio autoguiado o implementación/capacitación con socios (servicem8.com/pricing).',
 'Codelco SUCAL, SmartCheck, SubCheck, Fácil Control, Vigenty y Documental.cl: referencias locales para acreditación, contratistas, trabajadores, vehículos y matrices de cumplimiento.',
 'Los valores en pesos chilenos de este documento son una propuesta comercial de Nexo Klar, no una reproducción de tarifas de terceros; se deben validar con clientes piloto y costos de soporte/cloud reales.'
])

doc.save(str(OUT))
print(OUT)
