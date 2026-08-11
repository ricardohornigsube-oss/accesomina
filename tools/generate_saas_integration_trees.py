from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / 'docs' / 'Arboles_Integracion_Planes_SaaS_Nexo_Klar.docx'

INDIGO = '2A2A8C'
INDIGO_DEEP = '1A1A5E'
TEAL = '00CFC1'
TEAL_INK = '00706A'
INK = '141A20'
MUTED = '5D6B7A'
SURFACE = 'FFFFFF'
ALT = 'FBF9F5'
LINE = 'E3DED2'
GREEN = '1B7F4B'
AMBER = 'C77700'


def shade(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    element = OxmlElement('w:shd')
    element.set(qn('w:fill'), color)
    tc_pr.append(element)


def set_cell_border(cell, color=LINE, size='8'):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in('w:tcBorders')
    if borders is None:
        borders = OxmlElement('w:tcBorders')
        tc_pr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right'):
        node = borders.find(qn(f'w:{edge}'))
        if node is None:
            node = OxmlElement(f'w:{edge}')
            borders.append(node)
        node.set(qn('w:val'), 'single')
        node.set(qn('w:sz'), size)
        node.set(qn('w:color'), color)


def set_margins(cell, top=100, bottom=100, start=140, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = OxmlElement('w:tcMar')
    for key, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = OxmlElement(f'w:{key}')
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')
        margins.append(node)
    tc_pr.append(margins)


def add_run(paragraph, text, *, size=10, bold=False, color=INK):
    run = paragraph.add_run(text)
    run.font.name = 'Manrope'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Manrope')
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    return run


def title(doc, name, subtitle=''):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    add_run(p, name, size=25, bold=True, color=INDIGO_DEEP)
    if subtitle:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(16)
        add_run(p, subtitle, size=11, color=MUTED)


def heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(13)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    add_run(p, text, size=14, bold=True, color=INDIGO)


def body(doc, text, after=7):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    add_run(p, text, size=10, color=MUTED)


def card(doc, label, text, color=TEAL_INK):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, 'E6F2EB')
    set_cell_border(cell, 'BFE1D1')
    set_margins(cell, 130, 130, 170, 170)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    add_run(p, label.upper(), size=8.5, bold=True, color=color)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    add_run(p, text, size=9.2, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def tree_box(doc, title_text, nodes, accent=INDIGO):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, SURFACE)
    set_cell_border(cell)
    set_margins(cell, 150, 150, 180, 180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(7)
    add_run(p, title_text, size=11, bold=True, color=accent)
    for level, text, label in nodes:
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(0.2 + level * 0.28)
        p.paragraph_format.first_line_indent = Inches(-0.16)
        if level == 0:
            prefix = '● '
        elif level == 1:
            prefix = '├─ '
        else:
            prefix = '└─ '
        add_run(p, prefix, size=9.1, bold=True, color=accent)
        add_run(p, text, size=9.1, bold=level <= 1, color=INK)
        if label:
            add_run(p, f'  {label}', size=8.4, color=MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def integration_table(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [2200, 3450, 3260]
    headers = ['Origen', 'Conecta con', 'Resultado visible']
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        shade(cell, INDIGO)
        set_cell_border(cell, INDIGO)
        set_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        add_run(p, text, size=8, bold=True, color='FFFFFF')
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for index, row in enumerate(rows):
        cells = table.add_row().cells
        for cell_index, text in enumerate(row):
            cell = cells[cell_index]
            shade(cell, SURFACE if index % 2 == 0 else ALT)
            set_cell_border(cell)
            set_margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_run(p, text, size=8.25, color=INK)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def header_footer(section):
    p = section.header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p, 'NEXO KLAR SPA  |  Árboles de integración SaaS', size=8, bold=True, color=TEAL_INK)
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, 'Documento comercial de referencia · Nexo Klar · 2026', size=8, color=MUTED)


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(.7)
    section.bottom_margin = Inches(.65)
    section.left_margin = Inches(.72)
    section.right_margin = Inches(.72)
    header_footer(section)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(48)
    p.paragraph_format.space_after = Pt(8)
    add_run(p, 'NEXO KLAR', size=15, bold=True, color=TEAL_INK)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    add_run(p, 'Árboles de\nintegración SaaS', size=31, bold=True, color=INDIGO_DEEP)
    body(doc, 'Vista simple de cómo se relacionan los módulos de cada plan comercial, desde la empresa hasta el control de la operación.', after=12)
    card(doc, 'Cómo leer este documento', 'Cada árbol muestra únicamente los módulos visibles para el cliente según su plan. Las flechas y ramas indican qué información alimenta a cada control, alerta o reporte.')
    body(doc, f'Fecha de emisión: {date.today().strftime("%d/%m/%Y")}', after=0)

    doc.add_page_break()
    title(doc, 'Ruta común de integración', 'Todos los planes comparten una misma lógica de trabajo; aumenta la profundidad de control en cada nivel.')
    tree_box(doc, 'Ruta operativa base', [
        (0, 'Empresa', 'configuración, usuarios, permisos y catálogos'),
        (1, 'Cliente', 'contactos, antecedentes y condiciones'),
        (2, 'Contrato', 'alcance, vigencia, firmas y responsables'),
        (2, 'Orden de servicio', 'fechas, ubicación, responsable y estado'),
        (2, 'Personas y requisitos', 'turnos, EPP, formación, exámenes y aptitudes'),
        (2, 'Cumplimiento', 'documentos, habilitación, auditoría y alertas'),
        (1, 'Panel y reportes', 'lectura consolidada para decidir y actuar'),
    ])
    heading(doc, 'Qué cambia entre planes')
    body(doc, 'Esencial controla la relación comercial, las personas y el cumplimiento. Operación suma recursos físicos, terceros y coordinación diaria. Integral incorpora gobierno, análisis avanzado e integraciones preparadas para crecer.')
    integration_table(doc, [
        ['Cliente / contrato', 'Orden de servicio', 'Define el contexto comercial de la ejecución.'],
        ['Personas y requisitos', 'Orden de servicio + alertas', 'Indica si una persona está disponible, habilitada o restringida.'],
        ['Cumplimiento', 'Alertas + auditoría', 'Anticipa documentos vencidos, faltantes u observados.'],
        ['Toda la operación', 'Panel General + reportes', 'Entrega una visión común para administración y operación.'],
    ])

    doc.add_page_break()
    title(doc, 'Plan 1 · Nexo Klar Esencial', 'Orden y cumplimiento base para comenzar con una sola fuente de información.')
    tree_box(doc, 'Árbol de integración del plan Esencial', [
        (0, 'Empresa', 'usuarios, permisos, importación y configuración'),
        (1, 'Relación comercial', ''),
        (2, 'Clientes', 'datos y contactos'),
        (2, 'Contratos y firmas', 'vigencia, alcance y documentación'),
        (2, 'Órdenes de servicio', 'trabajo, fecha, responsable y estado'),
        (1, 'Capital humano', ''),
        (2, 'Personas', 'fijos, por proyecto, disponibles y restringidos'),
        (2, 'Turnos y asistencia', 'jornada, entrada, salida y HH'),
        (2, 'EPP / formación / exámenes / salud', 'requisitos por persona'),
        (1, 'Cumplimiento', ''),
        (2, 'Cumplimiento corporativo', 'antecedentes de la empresa'),
        (2, 'Habilitación del cliente y auditoría', 'revisión, estados y evidencia'),
        (1, 'Control', ''),
        (2, 'Panel General, Alertas y Reportes', 'prioridades y lectura consolidada'),
    ])
    heading(doc, 'Lectura simple para el cliente')
    body(doc, 'El administrador crea un cliente, registra el contrato y genera una orden. Luego vincula personas y valida que sus documentos, EPP, formación y exámenes permitan trabajar. El Panel General y Alertas muestran qué está listo y qué debe corregirse.')
    integration_table(doc, [
        ['Cliente', 'Contrato y orden de servicio', 'Cada trabajo conserva su contexto y vigencia.'],
        ['Persona', 'EPP, cursos, exámenes, salud y turnos', 'Ficha única y estado de habilitación.'],
        ['Documentación', 'Auditoría y alertas', 'Renovación sin perder historial.'],
        ['Datos operativos', 'Panel General y reportes base', 'Control simple de la empresa.'],
    ])

    doc.add_page_break()
    title(doc, 'Plan 2 · Nexo Klar Operación', 'Coordinación diaria de servicios, recursos, terceros y personal en terreno.')
    tree_box(doc, 'Árbol de integración del plan Operación', [
        (0, 'Todo el plan Esencial', 'clientes, contratos, órdenes, personas y cumplimiento'),
        (1, 'Centro Operativo', 'coordina la ejecución diaria'),
        (2, 'Personas por proyecto', 'convocatoria, confirmación y asignación'),
        (2, 'Comunicación y convocatorias', 'correo o WhatsApp con historial'),
        (2, 'Alojamientos y estadías', 'hotel, habitación, fechas y costo'),
        (2, 'Vehículos, activos y credenciales', 'recurso habilitado y responsable'),
        (1, 'Terceros y subcontratos', ''),
        (2, 'Contratos, personal y habilitaciones', 'cumplimiento por empresa externa'),
        (2, 'Evaluación de desempeño', 'continuidad y riesgo del tercero'),
        (1, 'Activos, equipos e inventario', ''),
        (2, 'Bodegas y existencias', 'stock, lote, ubicación y mínimos'),
        (2, 'Movimientos / préstamos / mantenimiento', 'responsable, costo e historial'),
        (1, 'Control de ejecución', ''),
        (2, 'Incidentes y no conformidades', 'evidencia, responsable y cierre'),
        (2, 'Alertas, Panel y Reportes operacionales', 'brechas y decisiones diarias'),
    ], TEAL_INK)
    heading(doc, 'Lectura simple para el cliente')
    body(doc, 'La orden de servicio se transforma en una vista operativa: reúne las personas asignadas, sus requisitos, el hotel, vehículos, activos, EPP, inventario y terceros. El Centro Operativo identifica brechas antes de comenzar y mantiene trazabilidad durante la ejecución.')
    integration_table(doc, [
        ['Orden de servicio', 'Personas, alojamiento, vehículos, credenciales y activos', 'Servicio preparado o brecha visible.'],
        ['Bodega / inventario', 'EPP, herramientas, materiales y préstamos', 'Disponibilidad real y responsable definido.'],
        ['Tercero', 'Personal, documentos, habilitación y desempeño', 'Empresa externa controlada antes de operar.'],
        ['Incidente', 'Acción, evidencia, alerta y reporte', 'Seguimiento hasta el cierre.'],
    ])

    doc.add_page_break()
    title(doc, 'Plan 3 · Nexo Klar Integral', 'Control ejecutivo, parámetros avanzados e integraciones para escalar.')
    tree_box(doc, 'Árbol de integración del plan Integral', [
        (0, 'Todo el plan Operación', 'operación diaria conectada'),
        (1, 'Gobierno y configuración avanzada', ''),
        (2, 'Matriz de requisitos', 'por cliente, contrato, cargo, orden o ubicación'),
        (2, 'Formularios y catálogos configurables', 'sin afectar a otras empresas'),
        (2, 'Usuarios, permisos y privacidad', 'acceso según responsabilidad'),
        (1, 'Analítica ejecutiva', ''),
        (2, 'Reportes por cliente, contrato y orden', 'cumplimiento, riesgos, dotación y recursos'),
        (2, 'Costos operativos', 'EPP, activos, alojamiento, mantenimiento e inventario'),
        (2, 'Indicadores y tendencias', 'alertas, desempeño y rentabilidad'),
        (1, 'Integraciones preparadas', ''),
        (2, 'Firma electrónica / OCR / mensajería', 'según proveedor y alcance contratado'),
        (2, 'ERP / API / portal de revisión', 'flujo externo controlado'),
        (1, 'Administración Nexo Klar', 'solo equipo interno: empresas, soporte, planes y respaldo'),
    ], INDIGO_DEEP)
    heading(doc, 'Lectura simple para el cliente')
    body(doc, 'Integral no cambia la ruta de trabajo: la profundiza. La empresa conserva los mismos clientes, contratos, órdenes y personas, pero puede definir requisitos propios, medir resultados, analizar costos e integrar proveedores externos sin perder trazabilidad ni separación de datos.')
    integration_table(doc, [
        ['Reglas por empresa', 'Matriz de requisitos y formularios', 'Configuración adaptable sin mezclar organizaciones.'],
        ['Datos operativos', 'Analítica ejecutiva y costos', 'Decisiones por cliente, contrato u orden.'],
        ['Documentos y firmas', 'OCR, firma y portal externo', 'Flujo preparado para integración controlada.'],
        ['Administración Nexo Klar', 'Clientes SaaS, soporte y respaldos', 'Operación central sin acceso cruzado a datos.'],
    ])

    doc.add_page_break()
    title(doc, 'Servicios adicionales', 'Se conectan a cualquiera de los tres planes sin alterar la lógica principal.')
    tree_box(doc, 'Libro de obra digital', [
        (0, 'Cliente', 'contexto de la relación'),
        (1, 'Contrato', 'alcance y responsable'),
        (1, 'Orden de servicio', 'ejecución específica'),
        (2, 'Anotación correlativa', 'avance, instrucción, acuerdo, consulta o incidente'),
        (2, 'Compromiso y evidencia', 'responsable, plazo, adjunto e historial'),
        (2, 'Firma o revisión', 'correo, WhatsApp o proveedor integrado'),
        (1, 'Alertas, Auditoría e Incidentes', 'seguimiento y cierre'),
    ], AMBER)
    tree_box(doc, 'Prospectos y oportunidades', [
        (0, 'Oportunidad', 'empresa, contacto, monto y etapa'),
        (1, 'Bitácora comercial', 'llamadas, reuniones, archivos y acuerdos'),
        (1, 'Próxima acción', 'responsable, fecha y proyección'),
        (1, 'Conversión', ''),
        (2, 'Cliente', 'cuando se formaliza la relación'),
        (2, 'Contrato', 'cuando se adjudica el acuerdo'),
        (2, 'Orden de servicio', 'cuando comienza la ejecución'),
    ], GREEN)
    heading(doc, 'Regla de negocio')
    body(doc, 'Libro de obra digital agrega formalidad a la ejecución de servicios. Prospectos y oportunidades agrega trazabilidad comercial antes de que exista un cliente o contrato. Ambos servicios reutilizan las entidades existentes y no crean silos de información.')
    card(doc, 'Resumen comercial', 'Nexo Klar Esencial ordena. Nexo Klar Operación ejecuta. Nexo Klar Integral dirige y escala. Los adicionales amplían la trazabilidad comercial y formal de cada servicio.')

    doc.save(OUT)
    print(OUT)


if __name__ == '__main__':
    build()
