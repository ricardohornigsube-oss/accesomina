#!/usr/bin/env python3
"""Generate the Nexo Klar data migration and onboarding proposal."""
from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / 'docs' / 'Nexo_Klar_Migracion_Historica_y_Carga_Inteligente.docx'
FLOW = ROOT / 'docs' / '_nexo_klar_migration_flow.png'
NAVY='141A20'; INDIGO='2A2A8C'; TEAL='00706A'; MUTED='5D6B7A'; LINE='D9DCE8'; PALE='F0F0FA'; ALT='FBF9F5'; CONTENT=6.6

def rgb(v): return RGBColor.from_string(v)
def shade(c, color):
    tcPr=c._tc.get_or_add_tcPr(); el=OxmlElement('w:shd'); el.set(qn('w:fill'),color); tcPr.append(el)
def border(c,color=LINE,size='8'):
    tcPr=c._tc.get_or_add_tcPr();b=tcPr.first_child_found_in('w:tcBorders')
    if b is None:b=OxmlElement('w:tcBorders');tcPr.append(b)
    for edge in ('top','left','bottom','right'):
        e=OxmlElement('w:'+edge);e.set(qn('w:val'),'single');e.set(qn('w:sz'),size);e.set(qn('w:color'),color);b.append(e)
def padding(c):
    tcPr=c._tc.get_or_add_tcPr();mar=OxmlElement('w:tcMar')
    for side,val in [('top','90'),('start','120'),('bottom','90'),('end','120')]:
        e=OxmlElement('w:'+side);e.set(qn('w:w'),val);e.set(qn('w:type'),'dxa');mar.append(e)
    tcPr.append(mar)
def no_split(row): row._tr.get_or_add_trPr().append(OxmlElement('w:cantSplit'))
def cell(c,text,bold=False,color=NAVY,size=9):
    c.text='';p=c.paragraphs[0];p.paragraph_format.space_before=Pt(1);p.paragraph_format.space_after=Pt(1)
    r=p.add_run(text);r.bold=bold;r.font.name='Arial';r.font.size=Pt(size);r.font.color.rgb=rgb(color);c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
def title(doc,text,sub=None):
    p=doc.add_paragraph();p.paragraph_format.space_before=Pt(10);p.paragraph_format.space_after=Pt(5)
    r=p.add_run(text);r.bold=True;r.font.name='Arial';r.font.size=Pt(24);r.font.color.rgb=rgb(NAVY)
    if sub:
        p=doc.add_paragraph();p.paragraph_format.space_after=Pt(12);r=p.add_run(sub);r.font.name='Arial';r.font.size=Pt(11);r.font.color.rgb=rgb(MUTED)
def body(doc,text):
    p=doc.add_paragraph();p.paragraph_format.line_spacing=1.13;p.paragraph_format.space_after=Pt(7)
    r=p.add_run(text);r.font.name='Arial';r.font.size=Pt(10);r.font.color.rgb=rgb(NAVY)
def bullets(doc,items):
    for item in items:
        p=doc.add_paragraph(style='List Bullet');p.paragraph_format.left_indent=Inches(.25);p.paragraph_format.space_after=Pt(3)
        r=p.add_run(item);r.font.name='Arial';r.font.size=Pt(9.5);r.font.color.rgb=rgb(NAVY)
def table(doc,headers,rows,widths):
    factor=CONTENT/sum(widths); widths=[w*factor for w in widths]
    t=doc.add_table(rows=1,cols=len(headers));t.alignment=WD_TABLE_ALIGNMENT.CENTER;t.autofit=False
    for i,h in enumerate(headers):
        t.columns[i].width=Inches(widths[i]);c=t.rows[0].cells[i];c.width=Inches(widths[i]);shade(c,INDIGO);border(c,INDIGO);padding(c);cell(c,h,True,'FFFFFF',8.5)
    for idx,row in enumerate(rows):
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].width=Inches(widths[i]);shade(cells[i],ALT if idx%2==0 else 'FFFFFF');border(cells[i]);padding(cells[i]);cell(cells[i],v,False,NAVY,8.5)
        no_split(t.rows[-1])
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
def callout(doc,label,text):
    t=doc.add_table(rows=1,cols=2);t.alignment=WD_TABLE_ALIGNMENT.CENTER;t.autofit=False
    for i,w in enumerate((1.3,5.3)):
        t.columns[i].width=Inches(w);t.cell(0,i).width=Inches(w);shade(t.cell(0,i),'EAFBF9');border(t.cell(0,i),'BDEDE8');padding(t.cell(0,i))
    cell(t.cell(0,0),label,True,TEAL,9);cell(t.cell(0,1),text,False,NAVY,9)
    doc.add_paragraph().paragraph_format.space_after=Pt(1)
def page(doc): doc.add_page_break()
def header_footer(sec):
    p=sec.header.paragraphs[0];p.alignment=WD_ALIGN_PARAGRAPH.RIGHT;r=p.add_run('NEXO KLAR SPA · MIGRACIÓN Y CARGA INTELIGENTE');r.font.name='Arial';r.font.size=Pt(7.5);r.font.color.rgb=rgb(MUTED)
    p=sec.footer.paragraphs[0];p.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=p.add_run('Documento comercial y técnico · Confidencial');r.font.name='Arial';r.font.size=Pt(7.5);r.font.color.rgb=rgb(MUTED)
def flow():
    im=Image.new('RGB',(1600,700),'#FBF9F5');d=ImageDraw.Draw(im)
    bold=ImageFont.truetype('/System/Library/Fonts/Supplemental/Arial Bold.ttf',30);reg=ImageFont.truetype('/System/Library/Fonts/Supplemental/Arial.ttf',19)
    d.text((65,45),'Fábrica de migración Nexo Klar',font=bold,fill='#141A20')
    boxes=[('1. Recibir','Excel, Word, PDF, fotos\ny carpetas'),('2. Proteger','Carga privada, antivirus\ny registro de lote'),('3. Leer','OCR, extracción de RUT,\nfechas y entidades'),('4. Ordenar','Clasificación, mapeo\ny deduplicación'),('5. Validar','Reglas, excepciones\ny revisión humana'),('6. Publicar','Datos trazables en\nNexo Klar')]
    x=55;y=210;bw=230;bh=150;gap=25
    for i,(a,b) in enumerate(boxes):
        xx=x+i*(bw+gap);d.rounded_rectangle((xx,y,xx+bw,y+bh),22,fill='#FFFFFF',outline='#2A2A8C',width=4)
        d.text((xx+18,y+25),a,font=bold,fill='#141A20');d.multiline_text((xx+18,y+72),b,font=reg,fill='#5D6B7A',spacing=4)
        if i<5:
            d.line((xx+bw,y+75,xx+bw+gap-6,y+75),fill='#00AFA5',width=5);d.polygon([(xx+bw+gap-6,y+75),(xx+bw+gap-22,y+66),(xx+bw+gap-22,y+84)],fill='#00AFA5')
    d.rounded_rectangle((125,470,1475,595),18,fill='#EEF6FF',outline='#B7C4E8',width=2)
    d.text((165,500),'Principio de control: ningún dato o documento se publica sin reglas, evidencia de origen y aprobación humana.',font=bold,fill='#2A2A8C')
    d.text((165,545),'Resultado: historia operacional buscable, datos activos confiables y archivos originales protegidos por empresa.',font=reg,fill='#5D6B7A')
    im.save(FLOW)

flow();doc=Document();sec=doc.sections[0];sec.top_margin=Inches(.65);sec.bottom_margin=Inches(.65);sec.left_margin=Inches(.8);sec.right_margin=Inches(.8);header_footer(sec)

doc.add_paragraph().paragraph_format.space_after=Pt(45)
p=doc.add_paragraph();r=p.add_run('NEXO KLAR SPA');r.bold=True;r.font.name='Arial';r.font.size=Pt(12);r.font.color.rgb=rgb(TEAL)
p=doc.add_paragraph();r=p.add_run('Migración Histórica\ny Carga Inteligente');r.bold=True;r.font.name='Arial';r.font.size=Pt(29);r.font.color.rgb=rgb(NAVY)
p=doc.add_paragraph();p.paragraph_format.space_after=Pt(18);r=p.add_run('Propuesta profesional para incorporar cinco años de datos operacionales, documentos y evidencias sin detener la operación.');r.font.name='Arial';r.font.size=Pt(13);r.font.color.rgb=rgb(MUTED)
callout(doc,'Propuesta','El cliente no necesita llegar ordenado. Nexo Klar recibe, protege, clasifica, valida y publica su información de forma gradual, segura y trazable.')
t=doc.add_table(rows=1,cols=1);t.alignment=WD_TABLE_ALIGNMENT.CENTER;t.autofit=False;t.columns[0].width=Inches(CONTENT);c=t.cell(0,0);c.width=Inches(CONTENT);shade(c,'FFFFFF');border(c,'D9DCE8','12');padding(c);p=c.paragraphs[0];p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.add_run().add_picture(str(FLOW),width=Inches(6.0))
page(doc)

title(doc,'1. Objetivo y alcance','Qué resuelve este servicio')
body(doc,'La migración histórica permite transformar planillas, documentos, fotografías, correos y carpetas dispersas en información útil para Personas, Clientes, Contratos, Órdenes de Servicio, Activos, EPP, Vehículos, Terceros, Alertas, Auditoría y Reportes de Nexo Klar.')
table(doc,['Información a recuperar','Resultado dentro de Nexo Klar'],[
['Personas y relación laboral','Nombre, RUT, contacto, dirección, cargo, tipo de trabajador, asignaciones e historial.'],
['Documentos y cumplimiento','Contratos, anexos, identificación, cursos, exámenes, credenciales, F30/F30-1 y vencimientos.'],
['Operación y servicios','Clientes, contratos, órdenes, turnos, actividades, libro de obra, incidentes y evidencia.'],
['Recursos y activos','Vehículos, equipos, herramientas, inventario, EPP, bodega, entrega y devolución.'],
['Historia consultable','Documentos originales, origen, año, entidad asociada, estado de revisión y bitácora.']
],[2.05,4.55])
callout(doc,'Resultado','Operación activa lista primero; historia de cinco años disponible, buscable y recuperable de manera progresiva.')

title(doc,'2. Modelo de migración recomendado','Activo primero, historia después')
table(doc,['Ola','Prioridad','Contenido','Resultado'],[
['Ola 0','Diagnóstico','Inventario de fuentes, muestra de calidad, volumen, riesgos, responsables y plan de migración.','Alcance aprobado y presupuesto cerrado.'],
['Ola 1','Crítica','Personas activas, documentos vigentes, clientes, contratos, órdenes activas, recursos y vencimientos.','Empresa operando en Nexo Klar.'],
['Ola 2','Operativa','Últimos 12–24 meses: EPP, trabajos, turnos, vehículos, incidentes, cursos y exámenes.','Trazabilidad operativa reciente.'],
['Ola 3','Histórica','Años restantes: respaldos, contratos cerrados, evidencia, informes y archivos antiguos.','Archivo histórico indexado y consultable.']
],[.8,1.05,3.05,1.7])
page(doc)

title(doc,'3. Alternativas de carga','Una solución para cada nivel de orden')
table(doc,['Alternativa','Cómo funciona','Cuándo usarla','Ventaja'],[
['Carga estructurada','Mapeo de columnas desde Excel/CSV, validaciones y carga por lotes.','Datos tabulares relativamente ordenados.','Rápida y de menor costo.'],
['Carga asistida','Nexo Klar limpia, transforma, relaciona y valida la información junto al cliente.','Planillas incompletas y archivos medianamente organizados.','Reduce trabajo del cliente.'],
['Rescate documental','Carpetas, PDFs, Word, fotos y correos pasan por OCR, clasificación y revisión humana.','Información dispersa o sin estructura.','Recupera evidencia sin exigir orden previo.'],
['Migración por unidad','Se migra cliente → contrato → orden de servicio → personas y recursos.','Muchas operaciones, zonas o contratos.','Genera valor desde la primera ola.'],
['Equipo dedicado','Jefe de proyecto, analista de datos y analista documental por hitos.','Alta complejidad, cinco años de historia y volumen corporativo.','Gobierno, velocidad y control de calidad.']
],[1.2,2.15,1.75,1.5])
callout(doc,'Decisión','Para cinco años de historia, la alternativa recomendada es Migración por Olas + Rescate Documental + Equipo Dedicado.')

title(doc,'4. Tecnología para transformar información desordenada','Automatización con control humano')
table(doc,['Tecnología','Qué hace','Uso concreto'],[
['Portal seguro de carga','Recibe lotes, carpetas ZIP, Excel, Word, PDF, fotos y enlaces cloud.','El cliente entrega información sin enviarla por correo o WhatsApp.'],
['OCR documental','Extrae texto, fechas, tablas, nombres, RUT, emisor y campos clave.','Clasificar contratos, exámenes, certificados, licencias y documentos de activos.'],
['Motor de clasificación','Propone tipo documental y entidad: persona, contrato, vehículo, tercero o servicio.','Reduce clasificación manual de miles de archivos.'],
['Reglas de calidad','Valida formato de RUT, fecha, duplicados, campos obligatorios y relación entre datos.','Evita publicar registros incompletos o repetidos.'],
['Bandeja de excepciones','Presenta solo casos dudosos, inconsistentes o sin relación clara.','Revisión rápida de un analista o responsable del cliente.'],
['Revisión humana','Aprueba, corrige o rechaza sugerencias antes de publicar.','Controla datos sensibles y documentación crítica.']
],[1.4,2.6,2.6])
page(doc)

title(doc,'5. Flujo operativo y de seguridad','Cómo se procesa un lote')
table(doc,['Paso','Control obligatorio','Salida'],[
['1. Recepción','Portal privado por empresa, inventario del lote y responsable de carga.','Lote identificado y trazable.'],
['2. Cuarentena','Antivirus, validación de formato, tamaño y tipo de archivo.','Archivos seguros para procesar.'],
['3. Extracción','OCR, lectura de metadatos, RUT, nombres, fechas y texto.','Datos candidatos y score de confianza.'],
['4. Normalización','Mapeo de Excel, estandarización de teléfono, fecha, región, comuna y nombres.','Datos comparables y reutilizables.'],
['5. Duplicados','Comparación por RUT, patente, contrato, correo, nombre y hash de archivo.','Registro nuevo, actualización o excepción.'],
['6. Validación','Reglas de negocio, matriz de requisitos y revisión humana.','Aprobado, observado, rechazado o pendiente.'],
['7. Publicación','Asociación con entidad correcta y bitácora de lote.','Datos disponibles en módulos y reportes.']
],[.8,3.25,2.55])
bullets(doc,[
'Todo archivo conserva su original, origen, fecha de ingreso, lote y usuario responsable.',
'Los datos se separan por empresa; un cliente nunca visualiza ni procesa información de otra organización.',
'Cada lote debe poder revertirse antes de su aceptación definitiva y quedar respaldado para auditoría.'
])
callout(doc,'Principio','La automatización acelera la lectura y ordenamiento; la publicación de información sensible requiere confirmación humana.')

title(doc,'6. Experiencia para el cliente','Carga simple, controlada y visible')
table(doc,['Pantalla propuesta','Función para el cliente'],[
['Centro de migración','Ver avance por lote, fuente, año, módulo, responsable, calidad y estado.'],
['Asistente de Excel','Relacionar columnas originales con campos Nexo Klar y guardar una plantilla reutilizable.'],
['Carga masiva de documentos','Arrastrar carpetas o ZIP, elegir contexto y dejar que el sistema proponga clasificación.'],
['Bandeja de revisión','Confirmar o corregir solo excepciones; no revisar documentos ya reconocidos con alta confianza.'],
['Informe de calidad','Ver duplicados, faltantes, archivos ilegibles, datos sin relación y campos incompletos.'],
['Aprobación de lote','Publicar un lote completo después de control interno y validación del cliente.']
],[2.1,4.5])
page(doc)

title(doc,'7. Paquetes comerciales sugeridos','Servicios claros y rentables')
table(doc,['Servicio','Incluye','Precio referencial neto CLP'],[
['Diagnóstico de migración','Muestra de fuentes, inventario, medición de volumen, plan por olas y presupuesto cerrado.','$750.000 a $1.500.000'],
['Inicio guiado','Plantillas, capacitación, carga de prueba y control de primer lote.','Incluido anual o $150.000'],
['Migración asistida','Hasta 50 personas, 2 clientes, 3 contratos/órdenes y documentos vigentes.','$490.000'],
['Migración operacional','Hasta 150 personas, 5 clientes, 10 contratos/órdenes, recursos, documentos y QA.','$990.000'],
['Rescate histórico','OCR, clasificación, indexación, depuración, carga progresiva y control de calidad.','Desde $2.500.000'],
['Migración corporativa 5 años','Equipo dedicado, datos maestros, documentos, recursos, revisión, publicación y capacitación.','Desde $8.000.000 a $20.000.000+']
],[1.55,3.55,1.5])
body(doc,'El precio final se calcula por volumen real: personas, archivos, años, formatos, calidad de origen, documentos sensibles, necesidad de OCR, revisión humana, integraciones y velocidad requerida. No se recomienda prometer migración total gratuita.')
callout(doc,'Oferta de cierre','“Comenzamos con su operación activa y recuperamos su historia por olas. Usted no necesita llegar ordenado para empezar.”')

title(doc,'8. Gobierno del proyecto','Responsabilidades y aceptación')
table(doc,['Nexo Klar','Cliente'],[
['Define modelo de datos, seguridad, plantillas, reglas, lotes, control de calidad, informe de errores y publicación.','Entrega fuentes autorizadas, designa responsable, valida reglas y aprueba lotes antes de producción.'],
['Protege originales, mantiene bitácora, administra excepciones y propone mejoras de calidad.','Aclara datos ambiguos, prioriza la información crítica y confirma relación de documentos sensibles.'],
['Entrega tablero semanal: recibidos, procesados, aprobados, observados, duplicados y pendientes.','Participa en revisión semanal y acepta cada hito de migración.']
],[3.3,3.3])
bullets(doc,[
'Hito 1: diagnóstico y alcance aprobado.',
'Hito 2: datos maestros y operación activa publicados.',
'Hito 3: historia operativa reciente validada.',
'Hito 4: archivo histórico indexado y proyecto cerrado.'
])
title(doc,'9. Roadmap técnico recomendado','Implementación progresiva sin sobreprometer')
table(doc,['Fase','Capacidad','Estado recomendado'],[
['Fase 1','Plantillas CSV, carga por lote, validación, deduplicación, almacenamiento privado y bitácora.','Necesaria antes de vender migración administrada.'],
['Fase 2','Portal de carga, bandeja de excepciones, informes de calidad y publicación controlada.','Prioridad alta para experiencia de cliente.'],
['Fase 3','OCR para PDFs e imágenes, detección de fechas, clasificación por tipo y score de confianza.','Activar mediante AWS Textract u otro proveedor.'],
['Fase 4','Clasificadores específicos de Nexo Klar, validación de emisor, QR, firma y matriz de requisitos.','Desarrollar con datos reales y validación humana.'],
['Fase 5','Conectores a correo, Drive, OneDrive, SharePoint, ERP, firma, WhatsApp y acreditación externa.','Cotizar por alcance y proveedor.']
],[.9,3.25,2.45])
callout(doc,'Tecnología','Para AWS, se recomienda S3 privado con enlaces temporales de carga, análisis OCR con Textract, procesamiento por colas y servicios de validación; cada tenant mantiene aislamiento lógico y auditoría.')

title(doc,'10. Referencias y propuesta de valor','Cómo comunicarlo')
body(doc,'AWS documenta el uso de enlaces temporales de S3 para cargas sin exponer credenciales y mecanismos de integridad. Amazon Textract ofrece OCR, extracción de formularios, tablas y scores de confianza; Azure Document Intelligence y Google Document AI son alternativas equivalentes según el entorno cloud del cliente.')
bullets(doc,[
'AWS S3 Presigned URLs: carga temporal y controlada de archivos.',
'Amazon Textract: OCR, formularios, tablas, extracción de información y niveles de confianza.',
'Azure Document Intelligence y Google Document AI: alternativas para OCR, clasificación y extracción de campos.',
'AWS Prescriptive Guidance: diagnóstico, priorización, validación, migración por olas y gobierno del proceso.'
])
callout(doc,'Mensaje comercial final','Nexo Klar transforma historia dispersa en operación visible. Protegemos los originales, ordenamos los datos, validamos la calidad y entregamos una plataforma lista para controlar el presente y responder por el pasado.')
doc.save(OUT);print(OUT)
