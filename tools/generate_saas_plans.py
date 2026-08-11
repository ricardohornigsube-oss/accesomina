from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / 'docs' / 'Propuesta_Planes_SaaS_Nexo_Klar.docx'

INDIGO = '2A2A8C'
INDIGO_DEEP = '1A1A5E'
TEAL = '00CFC1'
TEAL_INK = '00706A'
INK = '141A20'
MUTED = '5D6B7A'
BASE = 'F4EFE3'
SURFACE = 'FFFFFF'
ALT = 'FBF9F5'
LINE = 'E3DED2'
GREEN = '1B7F4B'
AMBER = 'C77700'


def shade(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    element = OxmlElement('w:shd')
    element.set(qn('w:fill'), color)
    tc_pr.append(element)


def set_cell_width(cell, value):
    tc_pr = cell._tc.get_or_add_tcPr()
    width = tc_pr.find(qn('w:tcW'))
    if width is None:
        width = OxmlElement('w:tcW')
        tc_pr.append(width)
    width.set(qn('w:w'), str(value))
    width.set(qn('w:type'), 'dxa')


def set_cell_margins(cell, top=90, bottom=90, start=130, end=130):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = OxmlElement('w:tcMar')
    for key, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = OxmlElement(f'w:{key}')
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')
        margins.append(node)
    tc_pr.append(margins)


def set_cell_border(cell, color=LINE, size='8'):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in('w:tcBorders')
    if borders is None:
        borders = OxmlElement('w:tcBorders')
        tc_pr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right'):
        node = borders.find(qn(f'w:{edge}'))
        if node is None:
            node = OxmlElement(f'w:{edge}')
            borders.append(node)
        node.set(qn('w:val'), 'single')
        node.set(qn('w:sz'), size)
        node.set(qn('w:color'), color)


def run(paragraph, text, size=10, bold=False, color=INK, font='Manrope'):
    item = paragraph.add_run(text)
    item.font.name = font
    item._element.rPr.rFonts.set(qn('w:eastAsia'), font)
    item.font.size = Pt(size)
    item.font.bold = bold
    item.font.color.rgb = RGBColor.from_string(color)
    return item


def cell_text(cell, text, *, bold=False, color=INK, size=8.5, align=None):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    if align:
        p.alignment = align
    run(p, str(text), size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def title(doc, text, subtitle=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run(p, text, size=26, bold=True, color=INDIGO_DEEP)
    if subtitle:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(16)
        run(p, subtitle, size=11, color=MUTED)


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run(p, text, size=16 if level == 1 else 12, bold=True, color=INDIGO if level == 1 else INK)
    return p


def body(doc, text, *, after=7):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    run(p, text, size=10, color=MUTED)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.08
    run(p, text, size=9.3, color=MUTED)
    return p


def note(doc, label, text, color=TEAL_INK):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, 'E6F2EB')
    set_cell_border(cell, 'BFE1D1')
    set_cell_margins(cell, 130, 130, 170, 170)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    run(p, label.upper(), size=8.5, bold=True, color=color)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run(p, text, size=9.2, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def table(doc, headers, rows, widths):
    content_width = sum(widths)
    result = doc.add_table(rows=1, cols=len(headers))
    result.autofit = False
    result.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, value in enumerate(headers):
        cell = result.rows[0].cells[idx]
        set_cell_width(cell, widths[idx])
        set_cell_border(cell, INDIGO)
        shade(cell, INDIGO)
        cell_text(cell, value, bold=True, color='FFFFFF', size=8)
    for row_index, row in enumerate(rows):
        row_cells = result.add_row().cells
        for idx, value in enumerate(row):
            cell = row_cells[idx]
            set_cell_width(cell, widths[idx])
            set_cell_border(cell)
            shade(cell, SURFACE if row_index % 2 == 0 else ALT)
            cell_text(cell, value, size=8.25)
    result.width = Inches(content_width / 1440)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return result


def plan_card(doc, number, name, purpose, audience, modules, highlight_color):
    card = doc.add_table(rows=1, cols=1)
    card.autofit = False
    card.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = card.cell(0, 0)
    shade(cell, SURFACE)
    set_cell_border(cell, LINE)
    set_cell_margins(cell, 170, 150, 200, 200)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    run(p, f'PLAN {number}', size=9, bold=True, color=highlight_color)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run(p, name, size=17, bold=True, color=INDIGO_DEEP)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run(p, purpose, size=10.5, bold=True, color=INK)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(9)
    run(p, f'Ideal para: {audience}', size=9.2, color=MUTED)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    run(p, 'Módulos incluidos', size=9, bold=True, color=TEAL_INK)
    for item in modules:
        p = cell.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(1)
        run(p, item, size=8.8, color=MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def add_header_footer(section):
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run(header, 'NEXO KLAR SPA  |  Propuesta de planes SaaS', size=8, bold=True, color=TEAL_INK)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(footer, 'Documento comercial de referencia · Nexo Klar · 2026', size=8, color=MUTED)


essential = [
    'Panel General y Alertas.',
    'Gestión de personal por proyecto.',
    'Personas, turnos y asistencia.',
    'Protección personal / EPP.',
    'Formación, certificaciones, exámenes y salud ocupacional.',
    'Restringidos y control de habilitación básica.',
    'Clientes, contratos y firmas.',
    'Órdenes de servicio.',
    'Cumplimiento corporativo, habilitación del cliente y auditoría documental.',
    'Reportes y analítica base.',
    'Configuración, importación/exportación, usuarios, permisos, bitácora y privacidad.',
]

operation = [
    'Todo lo incluido en Nexo Klar Esencial.',
    'Centro Operativo para coordinar servicios y brechas.',
    'Comunicaciones y convocatorias.',
    'Vehículos, activos y equipos.',
    'Alojamientos y estadías.',
    'Credenciales de acceso.',
    'Incidentes y no conformidades.',
    'Terceros y subcontratos.',
    'Contratos y convenios de terceros.',
    'Personal de empresa de servicios.',
    'Habilitaciones, cumplimiento y evaluación de desempeño de terceros.',
    'Inventario, EPP, materiales, bodegas, movimientos, mantenimiento, asignaciones y préstamos.',
]

integral = [
    'Todo lo incluido en Nexo Klar Operación.',
    'Reportes y analítica ejecutiva avanzada por cliente, contrato, orden, persona y recurso.',
    'Control operativo de costos: alojamiento, EPP, activos, mantenimiento e inventario.',
    'Matrices configurables de requisitos, formación, salud y EPP por cargo u orden.',
    'Reglas de alerta, tableros e indicadores avanzados.',
    'Configuración avanzada por empresa y formularios configurables.',
    'Portal de revisión para clientes, cuando se habilite.',
    'Conectores para API, ERP, firma electrónica, OCR y mensajería, según alcance contratado.',
    'Administración de clientes: uso exclusivo del equipo interno de Nexo Klar.',
]


def build():
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(.7)
    section.bottom_margin = Inches(.65)
    section.left_margin = Inches(.72)
    section.right_margin = Inches(.72)
    add_header_footer(section)

    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(48)
    p.paragraph_format.space_after = Pt(9)
    run(p, 'NEXO KLAR', size=15, bold=True, color=TEAL_INK)
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    run(p, 'Propuesta de\nplanes SaaS', size=31, bold=True, color=INDIGO_DEEP)
    body(document, 'Una oferta clara y escalable para ordenar personas, servicios, cumplimiento, recursos y decisiones en una única plataforma.', after=12)
    note(document, 'Propuesta comercial', 'La visibilidad del menú se adapta al plan contratado: el cliente ve solo las herramientas que necesita para operar, reduciendo complejidad y acelerando la adopción.')
    body(document, f'Fecha de emisión: {date.today().strftime("%d/%m/%Y")} · Valores comerciales, usuarios incluidos e integraciones se definen en la cotización particular de cada cliente.', after=0)

    document.add_page_break()
    title(document, 'Una oferta simple de entender', 'Tres niveles de adopción para acompañar el crecimiento operacional de cada empresa.')
    table(document,
          ['Plan', 'Nombre comercial', 'Enfoque', 'Cliente objetivo'],
          [
              ['1', 'Nexo Klar Esencial', 'Orden y cumplimiento base.', 'Empresas que comienzan a centralizar personas, contratos, documentos y alertas.'],
              ['2', 'Nexo Klar Operación', 'Coordinación diaria en terreno.', 'Empresas con órdenes activas, equipos, recursos, terceros y personal distribuido.'],
              ['3', 'Nexo Klar Integral', 'Control, decisión y escalabilidad.', 'Empresas con múltiples áreas, clientes, contratos y necesidades de gestión avanzada.'],
          ], [650, 2000, 2250, 3650])
    heading(document, 'Cómo se presenta al cliente')
    bullet(document, 'Esencial ordena la información crítica y evita vencimientos o faltantes.')
    bullet(document, 'Operación conecta personas, recursos, terceros y órdenes de servicio para ejecutar mejor.')
    bullet(document, 'Integral entrega una visión ejecutiva, parámetros avanzados e integraciones para escalar.')
    note(document, 'Servicios adicionales', 'Libro de obra digital y Prospectos y oportunidades se ofrecen como complementos activables en cualquiera de los tres planes.', AMBER)

    document.add_page_break()
    title(document, 'Plan 1 · Nexo Klar Esencial', 'Ordena la información crítica y controla vencimientos.')
    plan_card(document, '1', 'Nexo Klar Esencial', '“Una base única para personas, clientes, contratos, documentos y alertas.”', 'Empresas de servicios que hoy trabajan con planillas, carpetas, correos y controles dispersos.', essential, INDIGO)

    document.add_page_break()
    title(document, 'Plan 2 · Nexo Klar Operación', 'Coordina personas, recursos y servicios desde una misma operación.')
    plan_card(document, '2', 'Nexo Klar Operación', '“La operación diaria conectada con personas, recursos, cumplimiento y servicios en terreno.”', 'Contratistas, mantenimiento, construcción, energía, logística, facilities, seguridad y servicios técnicos.', operation, TEAL_INK)

    document.add_page_break()
    title(document, 'Plan 3 · Nexo Klar Integral', 'Control total para una operación conectada, trazable y escalable.')
    plan_card(document, '3', 'Nexo Klar Integral', '“Información ejecutiva, reglas configurables e integraciones para administrar una operación más compleja.”', 'Empresas con varios clientes, contratos, órdenes, recursos, terceros y equipos que requieren análisis y crecimiento controlado.', integral, INDIGO_DEEP)

    document.add_page_break()
    title(document, 'Servicios adicionales', 'Se activan en cualquiera de los planes según la necesidad comercial y operacional.')
    table(document,
          ['Servicio', 'Qué resuelve', 'Principales capacidades', 'Industrias de mayor uso'],
          [
              ['Libro de obra digital', 'Formaliza acuerdos, instrucciones, avances, incidentes y compromisos de cada servicio.', 'Folios correlativos, anotaciones, evidencia, responsables, plazos, historial, firma por correo o WhatsApp y relación con cliente, contrato y orden.', 'Construcción, mantenimiento, servicios industriales, proyectos, energía, facilities y contratos de terreno.'],
              ['Prospectos y oportunidades', 'Ordena la gestión comercial antes de crear un cliente, contrato u orden de servicio.', 'Embudo, etapas, responsable, próxima acción, bitácora, contacto, montos, documentos y conversión a cliente/contrato/orden.', 'Empresas de servicios, mantenimiento, contratación, ingeniería, logística y gestión comercial B2B.'],
          ], [1600, 2250, 2950, 2200])
    heading(document, 'Implementaciones opcionales')
    for item in [
        'Carga y depuración inicial de información mediante plantillas masivas.',
        'Parametrización de requisitos por cliente, contrato, cargo, orden o ubicación.',
        'Capacitación y acompañamiento de puesta en marcha.',
        'Integración con correo, WhatsApp, firma electrónica, OCR, ERP u otros sistemas.',
        'Identidad visual personalizada por empresa, bajo alcance contratado.',
    ]:
        bullet(document, item)

    document.add_page_break()
    title(document, 'Matriz de visibilidad SaaS', 'La plataforma conserva una experiencia clara: cada plan revela únicamente las capacidades contratadas.')
    table(document,
          ['Grupo funcional', 'Esencial', 'Operación', 'Integral'],
          [
              ['Panel, alertas y control base', 'Incluido', 'Incluido', 'Incluido + indicadores avanzados'],
              ['Personas, formación, salud, EPP y restricciones', 'Incluido', 'Incluido', 'Incluido + matrices configurables'],
              ['Clientes, contratos y órdenes de servicio', 'Incluido', 'Incluido', 'Incluido + analítica avanzada'],
              ['Recursos: vehículos, alojamiento y credenciales', '—', 'Incluido', 'Incluido'],
              ['Terceros y subcontratos', '—', 'Incluido', 'Incluido + evaluación avanzada'],
              ['Activos, equipos e inventario', '—', 'Incluido', 'Incluido + costos y conectores'],
              ['Cumplimiento y auditoría', 'Incluido', 'Incluido', 'Incluido + reglas y conectores'],
              ['Reportes y analítica', 'Base', 'Operacional', 'Ejecutiva avanzada'],
              ['Integraciones, formularios y portal cliente', '—', '—', 'Según alcance contratado'],
              ['Libro de obra digital', 'Adicional', 'Adicional', 'Adicional'],
              ['Prospectos y oportunidades', 'Adicional', 'Adicional', 'Adicional'],
          ], [3300, 1800, 1800, 2100])
    heading(document, 'Regla comercial recomendada')
    body(document, 'El cliente puede crecer de plan sin perder datos ni historial. Los servicios adicionales se activan de forma independiente; de esta manera, una empresa puede comenzar con control documental y luego incorporar operación, activos, Libro de obra o gestión comercial cuando su proceso lo requiera.')
    note(document, 'Resumen', 'Nexo Klar Esencial ordena. Nexo Klar Operación ejecuta. Nexo Klar Integral dirige y escala.', TEAL_INK)

    document.save(OUT)
    print(OUT)


if __name__ == '__main__':
    build()
