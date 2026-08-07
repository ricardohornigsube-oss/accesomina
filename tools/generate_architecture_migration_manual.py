from pathlib import Path
from datetime import date
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont

ROOT=Path(__file__).resolve().parents[1]
OUT=ROOT/'docs'/'Manual_Migracion_Arquitectura_Nexo_Klar_Cloud.docx'
ASSET=ROOT/'docs'/'_architecture_nexo_klar.png'
NAVY='14121F'; INDIGO='2A2A8C'; MAGENTA='E4006E'; MUTED='5F5B73'; LIGHT='F7F5FA'; LINE='DDD9E8'; GREEN='168A45'; ORANGE='F46F23'

def font(size=10,bold=False,color=NAVY,name='Aptos'):
    return {'name':name,'size':Pt(size),'bold':bold,'color':RGBColor.from_string(color)}
def shade(cell,color):
    tcPr=cell._tc.get_or_add_tcPr(); s=OxmlElement('w:shd'); s.set(qn('w:fill'),color); tcPr.append(s)
def margins(cell,top=90,start=110,bottom=90,end=110):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr(); m=tcPr.first_child_found_in('w:tcMar')
    if m is None: m=OxmlElement('w:tcMar');tcPr.append(m)
    for side,value in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        e=m.find(qn(f'w:{side}'))
        if e is None:e=OxmlElement(f'w:{side}');m.append(e)
        e.set(qn('w:w'),str(value));e.set(qn('w:type'),'dxa')
def cell_text(cell,text,bold=False,color=NAVY,size=9):
    cell.text=''; p=cell.paragraphs[0];p.paragraph_format.space_after=Pt(0);r=p.add_run(str(text));r.font.name='Aptos';r.font.size=Pt(size);r.font.bold=bold;r.font.color.rgb=RGBColor.from_string(color);cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER;margins(cell)
def set_repeat_table_header(row):
    trPr=row._tr.get_or_add_trPr();el=OxmlElement('w:tblHeader');el.set(qn('w:val'),'true');trPr.append(el)
def set_cell_width(cell,width):
    tcPr=cell._tc.get_or_add_tcPr(); tcW=tcPr.find(qn('w:tcW'))
    if tcW is None: tcW=OxmlElement('w:tcW');tcPr.append(tcW)
    tcW.set(qn('w:w'),str(width));tcW.set(qn('w:type'),'dxa')
def add_table(doc,headers,rows,widths):
    table=doc.add_table(rows=1, cols=len(headers));table.alignment=WD_TABLE_ALIGNMENT.CENTER;table.style='Table Grid';table.autofit=False
    for i,h in enumerate(headers):
        cell=table.rows[0].cells[i];set_cell_width(cell,widths[i]);shade(cell,INDIGO);cell_text(cell,h,True,'FFFFFF',8)
    set_repeat_table_header(table.rows[0])
    for n,row in enumerate(rows):
        cells=table.add_row().cells
        for i,val in enumerate(row):
            set_cell_width(cells[i],widths[i]);shade(cells[i],'FFFFFF' if n%2==0 else LIGHT);cell_text(cells[i],val,False,NAVY,8.2)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
    return table
def add_heading(doc,text,level=1):
    p=doc.add_paragraph(style=f'Heading {level}');p.paragraph_format.keep_with_next=True
    r=p.add_run(text);r.font.name='Aptos Display';r.font.color.rgb=RGBColor.from_string(INDIGO if level==1 else NAVY);return p
def add_body(doc,text,bold_lead=None):
    p=doc.add_paragraph();p.paragraph_format.space_after=Pt(7);p.paragraph_format.line_spacing=1.12
    if bold_lead and text.startswith(bold_lead):
        r=p.add_run(bold_lead);r.bold=True;r.font.color.rgb=RGBColor.from_string(NAVY);p.add_run(text[len(bold_lead):])
    else:p.add_run(text)
    for r in p.runs:r.font.name='Aptos';r.font.size=Pt(10);r.font.color.rgb=RGBColor.from_string(MUTED)
    return p
def bullets(doc,items):
    for item in items:
        p=doc.add_paragraph(style='List Bullet');p.paragraph_format.space_after=Pt(3);r=p.add_run(item);r.font.name='Aptos';r.font.size=Pt(9.5);r.font.color.rgb=RGBColor.from_string(MUTED)
def callout(doc,title,text,color=MAGENTA):
    table=doc.add_table(rows=1,cols=1);table.alignment=WD_TABLE_ALIGNMENT.CENTER;cell=table.cell(0,0);shade(cell,'FFF3F9');margins(cell,140,180,140,180)
    p=cell.paragraphs[0];p.paragraph_format.space_after=Pt(2);r=p.add_run(title.upper());r.bold=True;r.font.size=Pt(9);r.font.name='Aptos';r.font.color.rgb=RGBColor.from_string(color)
    p=cell.add_paragraph();p.paragraph_format.space_after=Pt(0);r=p.add_run(text);r.font.name='Aptos';r.font.size=Pt(9.5);r.font.color.rgb=RGBColor.from_string(NAVY)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
def page_break(doc): doc.add_page_break()

def diagram():
    im=Image.new('RGB',(1800,980),'white');d=ImageDraw.Draw(im)
    try:f_b=ImageFont.truetype('/System/Library/Fonts/Supplemental/Arial Bold.ttf',33);f=ImageFont.truetype('/System/Library/Fonts/Supplemental/Arial.ttf',24)
    except: f_b=f=ImageFont.load_default()
    boxes=[('Usuarios\nNavegador',70,350,300,490,'E4006E'),('HTTPS + ALB\nWAF / ACM',390,350,650,490,'2A2A8C'),('ECS Fargate\nNode.js + Express',750,320,1070,520,'2A2A8C'),('RDS PostgreSQL\nRLS + auditoría',1210,210,1510,360,'168A45'),('S3 privado\narchivos + versiones',1210,425,1510,575,'F46F23'),('Secrets Manager\ncredenciales',1210,640,1510,790,'5F5B73'),('CloudWatch\nlogs + alarmas',750,690,1070,840,'5F5B73')]
    for label,x1,y1,x2,y2,c in boxes:
        d.rounded_rectangle((x1,y1,x2,y2),radius=20,fill='#FFFFFF',outline='#'+c,width=5)
        lines=label.split('\n');yy=(y1+y2)//2-((len(lines)*35)//2)
        for line in lines:
            w=d.textbbox((0,0),line,font=f_b)[2];d.text(((x1+x2-w)//2,yy),line,font=f_b,fill='#14121F');yy+=38
    arrows=[(300,420,390,420),(650,420,750,420),(1070,375,1210,285),(1070,445,1210,500),(1070,500,1210,710),(910,520,910,690)]
    for x1,y1,x2,y2 in arrows:d.line((x1,y1,x2,y2),fill='#2A2A8C',width=5);d.polygon([(x2,y2),(x2-16,y2-10),(x2-16,y2+10)],fill='#2A2A8C')
    d.text((70,85),'Nexo Klar Cloud · arquitectura productiva multiempresa',font=f_b,fill='#14121F')
    d.text((70,140),'Cada solicitud autenticada recibe un contexto tenant_id; PostgreSQL aplica aislamiento RLS y la auditoría es append-only.',font=f,fill='#5F5B73')
    im.save(ASSET)

def build():
    diagram();doc=Document();sec=doc.sections[0];sec.top_margin=Inches(.72);sec.bottom_margin=Inches(.7);sec.left_margin=Inches(.78);sec.right_margin=Inches(.78)
    styles=doc.styles
    styles['Normal'].font.name='Aptos';styles['Normal'].font.size=Pt(10);styles['Normal'].font.color.rgb=RGBColor.from_string(MUTED)
    for name,size,color in [('Title',30,NAVY),('Subtitle',13,MUTED),('Heading 1',19,INDIGO),('Heading 2',14,NAVY),('Heading 3',11,INDIGO)]:
        st=styles[name];st.font.name='Aptos Display' if name!='Normal' else 'Aptos';st.font.size=Pt(size);st.font.color.rgb=RGBColor.from_string(color);st.font.bold=name!='Subtitle'
    # Header/footer
    header=sec.header.paragraphs[0];header.alignment=WD_ALIGN_PARAGRAPH.RIGHT;r=header.add_run('NEXO KLAR SPA  |  Manual técnico de migración');r.font.name='Aptos';r.font.size=Pt(8);r.font.color.rgb=RGBColor.from_string(MAGENTA)
    footer=sec.footer.paragraphs[0];footer.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=footer.add_run('Documento controlado · Arquitectura y migración a producción · Agosto 2026');r.font.name='Aptos';r.font.size=Pt(8);r.font.color.rgb=RGBColor.from_string(MUTED)
    # cover
    p=doc.add_paragraph();p.paragraph_format.space_before=Pt(90);p.paragraph_format.space_after=Pt(10);r=p.add_run('NEXO KLAR');r.font.name='Aptos Display';r.font.size=Pt(16);r.font.bold=True;r.font.color.rgb=RGBColor.from_string(MAGENTA)
    p=doc.add_paragraph();p.paragraph_format.space_after=Pt(12);r=p.add_run('Manual maestro de\narquitectura y migración cloud');r.font.name='Aptos Display';r.font.size=Pt(34);r.font.bold=True;r.font.color.rgb=RGBColor.from_string(NAVY)
    add_body(doc,'Guía integral para integrar, desplegar, operar y validar Nexo Klar en producción. Cubre la aplicación pública y privada, los módulos funcionales, APIs, datos, seguridad, AWS, migraciones, continuidad y pruebas de salida.')
    callout(doc,'Propósito','Entregar a un integrador una ruta única para mover el repositorio GitHub hacia una operación SaaS productiva, segura, multiempresa y trazable.',INDIGO)
    doc.add_paragraph('Versión de referencia: 7.7.0 · Arquitectura actualizada con Centro de control SaaS por empresa · Fecha: '+date.today().strftime('%d/%m/%Y')).runs[0].font.size=Pt(9)
    page_break(doc)
    add_heading(doc,'1. Cómo leer este manual')
    add_body(doc,'Este documento está diseñado para arquitectura, DevOps, desarrollo backend, seguridad, soporte y dirección de producto. La aplicación es una plataforma SaaS multiempresa; su fuente autoritativa en producción es PostgreSQL, no el navegador ni GitHub.')
    add_table(doc,['Capa','Responsabilidad','Fuente principal'],[
      ['Presentación','Sitio público, acceso y sitio privado','public/index.html y AccesoMina_v6.html'],
      ['Cliente cloud','Sesión, persistencia modular, API y carga de archivos','public/cloud-client.js'],
      ['API','Autenticación, reglas, permisos e integraciones','server/index.js y server/routes/'],
      ['Datos','Aislamiento, integridad, auditoría y estado','PostgreSQL 15+ / database/postgres/'],
      ['Archivos','Evidencia privada, versión y análisis','Amazon S3 privado + antivirus'],
      ['Operación','Despliegue, métricas, secretos y respaldos','infra/aws/ y AWS administrado']
    ],[1550,3150,3660])
    add_heading(doc,'2. Principios no negociables')
    bullets(doc,['Cada empresa tiene un tenant_id. El navegador no decide el tenant; la API lo obtiene desde la sesión.', 'Toda consulta operacional usa contexto de tenant y PostgreSQL aplica Row-Level Security forzado.', 'Los cambios se guardan por módulo con versión optimista. Un cambio desactualizado devuelve conflicto y debe recargarse.', 'La auditoría es append-only: guarda usuario, fecha, entidad, acción, valor anterior y nuevo.', 'Archivos, credenciales y secretos no se almacenan en GitHub ni como Base64 en estado operacional.', 'La demo file:// se mantiene para presentación; no es una modalidad productiva.'])
    page_break(doc)
    add_heading(doc,'3. Arquitectura de referencia')
    doc.add_picture(str(ASSET),width=Inches(6.85));doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    add_body(doc,'Flujo principal: el usuario navega mediante HTTPS. El Application Load Balancer entrega tráfico a ECS Fargate. Node.js valida sesión, CSRF, MFA y rol; abre una transacción con el tenant correcto. RDS aplica RLS e integridad. Los archivos se alojan en S3 privado y los secretos se obtienen desde Secrets Manager.')
    add_table(doc,['Componente','Función en producción','Configuración crítica'],[
      ['GitHub','Código, revisiones, pipeline e historial','main protegido, secretos solo en GitHub/AWS'],['ECR','Imágenes versionadas del contenedor','Etiquetas inmutables por commit'],['ECS Fargate','Ejecuta API Node.js sin servidores','Subred privada, autoscaling, rol IAM mínimo'],['ALB + ACM','HTTPS, healthcheck y enrutamiento','443, certificado válido, /api/health'],['RDS PostgreSQL','Datos multiempresa y auditoría','Multi-AZ, cifrado, backups, no público'],['S3 privado','Documentos, fotos y respaldos','Block Public Access, versionado, lifecycle'],['Secrets Manager','Claves y tokens','Rotación y permisos por rol'],['CloudWatch','Logs, métricas y alarmas','Alertas de error, caída y capacidad']
    ],[1600,3700,3060])
    page_break(doc)
    add_heading(doc,'4. Estado actual: legado, compatibilidad y destino productivo')
    add_table(doc,['Elemento','Estado actual','Uso futuro / decisión de integración'],[
      ['AccesoMina_v6.html','Aplicación local y demo, sincronizada con public/index.html','Mantener para demostración y QA visual local. No usar como almacenamiento productivo.'],['public/index.html','Frontend público y privado servido en cloud','Artefacto de producción; debe mantenerse idéntico a la versión local.'],['localStorage','Compatibilidad para demo file://','No usar como fuente de verdad en cloud.'],['tenant_state','Snapshot completo heredado','Compatibilidad, recuperación y primera migración.'],['tenant_module_state','Colecciones por módulo con versiones','Fuente operacional inmediata; permite concurrencia por módulo.'],['Tablas normalizadas','Workers, contratos, EPP, inventario, libro, etc.','Destino para APIs especializadas, analítica y escalamiento.'],['JSON / JSONB','Catálogos, formularios, evidencia, configuración y respaldos','Flexibilidad por empresa sin perder relaciones críticas.']
    ],[1700,2900,3760])
    callout(doc,'Decisión recomendada','Conservar arquitectura híbrida. PostgreSQL normalizado para claves, relaciones, inventario, auditoría y reportes; JSONB para configuraciones y formularios; JSON para respaldo e integración.',MAGENTA)
    add_heading(doc,'5. Modelo multiempresa y seguridad de datos')
    bullets(doc,['tenants identifica cada empresa, su RUT, administrador y estado.', 'app_users y user_sessions controlan accesos individuales, sesiones revocables, MFA y bloqueos.', 'tenant_settings, tenant_integrations, tenant_catalog_items y tenant_form_definitions permiten configuración aislada.', 'tenant_commercial_profiles y tenant_support_tickets controlan plan, onboarding, soporte, pago y ciclo de vida SaaS.', 'RLS usa app.current_tenant_id dentro de una transacción; las tablas críticas se crean con FORCE ROW LEVEL SECURITY.', 'El rol domian_admin administra empresas; client_admin administra solo su empresa.'])
    page_break(doc)
    add_heading(doc,'6. Mapa funcional completo')
    modules=[
      ('Centro de control','Panel General','Indicadores, alertas, tareas y acceso a prioridades operativas.'),('Centro de control','Alertas','Vencimientos, faltantes, agrupación por persona y carga directa de evidencia.'),('Centro de control','Gestión de personal por proyecto','Convocatoria, disponibilidad, contratación y habilitación por orden.'),('Centro de control','Centro Operativo','Brechas, disponibilidad, recursos, inventario, costos y control de ejecución.'),
      ('Capital humano','Personas','Ficha central: datos, documentos, cursos, exámenes, EPP, asignaciones, hotel y credenciales.'),('Capital humano','Trabajador fijo','Dotación estable, cargo, turno, centro de costo y cumplimiento.'),('Capital humano','Trabajador por proyecto','Personal temporal asociado a una orden o contrato.'),('Capital humano','Trabajador disponible','Pool habilitable para futuras convocatorias.'),('Capital humano','Restringidos','Personas con bloqueo operacional, documental o de seguridad.'),('Capital humano','Turnos y asistencia','Turnos, ingreso/salida, HH, dotación y jornadas.'),('Capital humano','Protección personal / EPP','Matriz por función, tallas, entregas, reposición, costo y respaldo.'),('Capital humano','Formación y certificaciones','Cursos obligatorios, certificaciones, vigencias y matrices.'),('Capital humano','Exámenes y aptitudes','Preocupacional, aptitud, vencimientos y evidencia.'),('Capital humano','Salud ocupacional','Protocolos por exposición, riesgo y control de datos sensibles.'),
      ('Relación comercial','Prospectos y oportunidades','Embudo, contactos, bitácora, documentos, montos y conversión.'),('Relación comercial','Clientes','Ficha, contactos, historial, contratos, órdenes y datos comerciales.'),('Relación comercial','Contratos y firmas','Plantillas, anexos, firma, vigencia, versiones y alertas.'),('Relación comercial','Órdenes de servicio','Planificación, presupuesto, personas, recursos, avance y estado operativo.'),
      ('Gestión operacional','Terceros y subcontratos','Empresa externa, cumplimiento, documentos, personal, evaluación y contratos.'),('Gestión operacional','Comunicaciones y convocatorias','WhatsApp/correo, grupos, consentimiento, respuesta y trazabilidad.'),('Gestión operacional','Vehículos, activos y equipos','Patente/serie, seguros, vencimientos, asignación, costos y mantenimiento.'),('Gestión operacional','Alojamientos y estadías','Hoteles, camas, tarifas, disponibilidad, estadías e historial.'),('Gestión operacional','Credenciales de acceso','Pases, QR, zonas, vencimientos y bloqueo por incumplimiento.'),('Gestión operacional','Incidentes y no conformidades','Registro, evidencia, investigación, CAPA, responsables y cierre.'),
      ('Activos, equipos e inventario','Maquinaria / equipos / herramientas','Fichas, códigos, ubicación, estado, mantenimiento y asignaciones.'),('Activos, equipos e inventario','Materiales, insumos y EPP','Stock, lotes, vida útil, mínimo, costo y reposición.'),('Activos, equipos e inventario','Bodegas y movimientos','Bodegas, traslados, recepción, devolución, daño y respaldo firmado.'),('Activos, equipos e inventario','Asignaciones y préstamos','Entrega a persona u orden, fecha prevista de devolución e historial.'),
      ('Cumplimiento y calidad','Cumplimiento corporativo','Documentos de empresa, versiones, vigencias y trazabilidad.'),('Cumplimiento y calidad','Habilitación del cliente','Matriz de requisitos, estado de aprobación y observaciones.'),('Cumplimiento y calidad','Auditoría','Flujo documental, evidencia, revisión humana y señales para OCR/QR/firma.'),
      ('Proyectos y negocios','Libro de obra','Folio correlativo, anotaciones, adjuntos, compromisos, aprobaciones y firma.'),('Decisión y gobierno','Reportes y analítica','Reportes por cliente, contrato, orden, persona, recursos, costos y cumplimiento.'),('Decisión y gobierno','Configuración de la empresa','Marca, tema, alertas, módulos, catálogos y especialidades.'),('Decisión y gobierno','Importar y exportar','CSV/JSON, plantillas, validación, duplicados y respaldo.'),('Decisión y gobierno','Usuarios y permisos','Usuarios, roles, MFA, restablecimiento y permisos por módulo.'),('Decisión y gobierno','Bitácora y privacidad','Auditoría inmutable, derechos de datos, consentimientos y conservación.'),('Administración Nexo Klar','Administración de clientes','Ficha 360 SaaS, planes, soporte, seguridad, onboarding y respaldo por tenant.')]
    add_table(doc,['Grupo','Módulo','Interacción principal'],modules,[1800,2450,4110])
    page_break(doc)
    add_heading(doc,'7. Flujos de datos e interacciones')
    add_heading(doc,'7.1 Flujo comercial y operativo')
    add_body(doc,'Prospecto → Cliente → Contrato → Orden de servicio → Personas y recursos → Ejecución → Libro de obra / incidentes → Auditoría → Reportes. Cada elemento puede conservar bitácora, documentos, responsables, fechas y evidencia.')
    add_table(doc,['Origen','Relación obligatoria / recomendada','Resultado'],[
      ['Cliente','Contratos y contactos','Contexto comercial, histórico y responsables.'],['Contrato','Cliente y órdenes de servicio','Vigencia, alcance, precio, firma y anexos.'],['Orden de servicio','Cliente, contrato, fechas y responsable','Centro de asignación de personas, hotel, vehículos, EPP y activos.'],['Persona','Tipo, RUT único, documentos y aptitud','Puede estar fija, por proyecto, disponible o restringida.'],['EPP / activos','Bodega, stock, persona u orden','Descuento de existencias, costo, préstamo y devolución.'],['Alojamiento','Hotel, habitación/cama, persona, fechas','Controla disponibilidad y libera cama al checkout.'],['Libro de obra','Cliente, contrato y orden','Evidencia formal, compromisos y firmas.'],['Reportes','Filtros por entidad y periodo','Decisión ejecutiva, cumplimiento y costos.']
    ],[1550,4200,2610])
    add_heading(doc,'7.2 Reglas de integridad ya soportadas')
    bullets(doc,['RUT de persona único por empresa, incluso con formatos distintos.', 'Patente o serie de vehículo/equipo única dentro de la empresa.', 'Número de contrato, asignaciones y documentos críticos no se duplican.', 'No se permiten referencias entre empresas, clientes, contratos u órdenes incompatibles.', 'Las fechas de contrato, documento, alojamiento, arriendo y vencimiento se validan.', 'Una cama no admite estadías superpuestas y se libera al terminar.', 'Las anotaciones firmadas o cerradas del Libro de obra no se modifican; se corrigen con nueva anotación.', 'Los cambios de módulos aplican versión optimista para prevenir sobrescritura concurrente.'])
    page_break(doc)
    add_heading(doc,'8. API y contratos de integración')
    add_body(doc,'Base URL: https://<dominio>/api. Las rutas, excepto health/readiness y callbacks autorizados, usan sesión HTTP-only, token CSRF, control de origen y MFA cuando está exigido.')
    api_rows=[
      ('/api/auth','POST register/login/logout, GET me, MFA setup/enable/disable, change-password','Autenticación, cuenta y doble factor.'),('/api/state','GET, PUT modules','Carga y persistencia versionada por módulo.'),('/api/users','GET/POST/PATCH, reset-password, reset-mfa','Usuarios y permisos por tenant.'),('/api/tenants','GET/POST/PATCH/DELETE, control, tickets, export-backup','Administración global SaaS por Nexo Klar.'),('/api/files','POST, GET /:id, DELETE /:id','Carga y descarga autorizada de archivos privados.'),('/api/data-transfer','Export, template, import por tipo','Respaldo y carga masiva CSV/JSON.'),('/api/settings','GET/PUT, integrations/:provider','Marca, catálogos, alertas, módulos e integraciones por empresa.'),('/api/operations','Overview, documents, notifications, portal, forms, retention','Motor documental, notificaciones, portales y formularios.'),('/api/work-books','Entries, signatures, approvals','Libro de obra, firma y aprobación.'),('/api/integrations','Email, WhatsApp, signature, ERP, accreditation, events','Conectores externos y trazabilidad.'),('/api/privacy','Activities, requests, consents, incidents','Gobierno y protección de datos.'),('/api/audit','GET','Lectura autorizada de bitácora inmutable.'),('/api/health / ready / metrics','GET','Monitoreo, readiness y métricas protegidas.')]
    add_table(doc,['Ruta','Operaciones','Finalidad'],api_rows,[2100,3400,2860])
    add_heading(doc,'8.1 Integraciones externas')
    bullets(doc,['SMTP: correo transaccional y notificaciones. Configurar remitente, SPF/DKIM/DMARC y credenciales.', 'Meta WhatsApp Cloud API: mensajes individuales y grupales según consentimiento y plantillas aprobadas.', 'Firma electrónica: endpoint configurable y callback con SIGNATURE_WEBHOOK_SECRET.', 'OCR/IA documental: DOCUMENT_AI_API_URL/TOKEN; la aprobación humana sigue siendo obligatoria.', 'ERP y acreditación: webhooks configurables por tenant. No automatizar portales sin API y autorización formal.'])
    page_break(doc)
    add_heading(doc,'9. Base de datos y migraciones')
    add_body(doc,'PostgreSQL 15+ es obligatorio. Las migraciones se ejecutan en orden numérico desde database/postgres y se registran en schema_migrations. La versión actual incluye 001 a 021.')
    migrations=[
      ('001-003','Esquema base, tenant, personas, contratos, RLS y seed administrador.'),('004-008','Runtime cloud, sesiones, archivos, estado modular, configuraciones y gobierno Domian.'),('009-011','Tipos de operación, privacidad y deduplicación de archivos.'),('012-016','Centro Operativo, producción, MFA, modelo normalizado, despacho, finanzas e inventario.'),('017-020','Libro de obra, solicitud de firma, identidad visual y aprobaciones gobernadas.'),('021','Centro de control SaaS: plan, onboarding, soporte, ciclo de vida y respaldos por empresa.')]
    add_table(doc,['Migraciones','Resultado'],migrations,[2300,6060])
    add_heading(doc,'9.1 Tablas críticas')
    add_table(doc,['Dominio','Tablas principales','Uso'],[
      ['Identidad','tenants, app_users, user_sessions','Empresa, usuarios, roles, MFA y sesiones.'],['Estado modular','tenant_state, tenant_module_state','Compatibilidad y colecciones versionadas.'],['Configuración','tenant_settings, tenant_integrations, tenant_catalog_items, tenant_form_definitions','Parámetros y secretos por empresa.'],['Personas','workers, worker_documents, occupational_exams, courses','Identidad, antecedentes, aptitud y formación.'],['Operación','commercial_contracts, maintenance_projects, worker_assignments, vehicles, lodging','Contrato, órdenes, personas, vehículos y estadías.'],['Recursos','epp_items, epp_deliveries, inventory_items, inventory_movements','EPP, activos, stock y movimientos.'],['Gobierno','audit_log, operational_events, integration_events','Historial, monitoreo e integraciones.'],['Libro de obra','work_books, work_book_entries, history, approvals, signature_requests','Evidencia formal, revisión y firma.'],['SaaS','tenant_commercial_profiles, tenant_support_tickets','Plan, pago, soporte y ciclo de vida.']
    ],[1650,3900,2810])
    page_break(doc)
    add_heading(doc,'10. Manual de migración: GitHub a AWS')
    add_heading(doc,'10.1 Pre-requisitos')
    bullets(doc,['Cuenta AWS separada o al menos ambientes development, staging y production.', 'Dominio productivo, por ejemplo nexo.nexoklar.cl, y zona DNS controlada.', 'RDS PostgreSQL privado, S3 privado, ECS Fargate, ECR, ALB, ACM, Secrets Manager y CloudWatch.', 'Proveedor de correo, WhatsApp, antivirus, firma/OCR/ERP si corresponden.', 'Acceso GitHub con permisos de lectura y despliegue, sin secretos en el repositorio.'])
    add_heading(doc,'10.2 Orden de despliegue recomendado')
    steps=['Crear red: VPC, subredes privadas para ECS/RDS, subred pública solo para ALB y NAT controlado.', 'Crear RDS PostgreSQL 15+, cifrado, Multi-AZ, backups, monitoreo y security group exclusivo para ECS.', 'Crear bucket S3 con Block Public Access, SSE, versionado, lifecycle y rol IAM exclusivo de la tarea ECS.', 'Crear secretos de producción en Secrets Manager; no copiar valores en archivos .env ni imágenes.', 'Crear repositorio ECR y pipeline que ejecute pruebas antes de publicar imagen.', 'Crear definición ECS usando infra/aws/ecs-task-definition.json como plantilla y rol IAM mínimo.', 'Crear ALB HTTPS con certificado ACM; configurar healthcheck /api/health.', 'Ejecutar migraciones como tarea controlada. No habilitar RUN_MIGRATIONS_ON_START en múltiples réplicas.', 'Ejecutar seed inicial solo en ambiente nuevo, configurando la cuenta administradora Nexo Klar.', 'Configurar APP_ORIGIN, cookie segura, MFA obligatorio y registro público cerrado.', 'Desplegar servicio ECS y verificar /api/health y /api/ready.', 'Realizar pruebas de aislamiento, carga de archivos, MFA, respaldo/restauración y alertas antes de abrir clientes.']
    for i,s in enumerate(steps,1):
        p=doc.add_paragraph(style='List Number');p.paragraph_format.space_after=Pt(4);r=p.add_run(s);r.font.name='Aptos';r.font.size=Pt(9.5);r.font.color.rgb=RGBColor.from_string(MUTED)
    add_heading(doc,'10.3 Variables obligatorias')
    add_table(doc,['Variable','Uso','Producción'],[
      ['NODE_ENV','Modo de ejecución','production'],['APP_ORIGIN','Origen permitido','https://<dominio>'],['DATABASE_URL','Conexión RDS','Secret Manager'],['TENANT_SECRET_KEY','Cifrado de integraciones','Secret Manager; 32+ caracteres'],['MFA_REQUIRED','Exige doble factor','true'],['COOKIE_SECURE','Cookie HTTPS','true'],['FILE_STORAGE','Proveedor de archivos','s3'],['AWS_S3_BUCKET / AWS_REGION','Bucket privado','Valores reales'],['VIRUS_SCAN_API_URL / TOKEN','Escaneo de archivos','Proveedor contratado'],['METRICS_TOKEN','Protege métricas','Secret Manager'],['SIGNATURE_* / SMTP_* / WHATSAPP_*','Integraciones externas','Solo cuando estén contratadas']
    ],[2400,3500,2460])
    page_break(doc)
    add_heading(doc,'11. Seguridad, privacidad y continuidad')
    add_table(doc,['Control','Implementación actual','Acción del integrador'],[
      ['Autenticación','Contraseña scrypt, sal único, cambio de clave temporal','Forzar política, MFA y canal seguro de entrega.'],['Sesiones','Cookie HttpOnly, Secure, SameSite Strict, revocación','Configurar HTTPS y TTL; supervisar sesiones anómalas.'],['Autorización','Roles y permisos de módulo','Probar mínimo privilegio por rol.'],['Aislamiento','tenant_id + RLS + contexto transaccional','No conceder bypass RLS a la aplicación.'],['Auditoría','audit_log append-only','Retener, respaldar y restringir lectura.'],['Archivos','SHA-256, máximo 25MB, estado antivirus','S3 privado, escaneo real y URLs temporales.'],['Secretos','AES-256-GCM por tenant para integraciones','Guardar TENANT_SECRET_KEY en Secrets Manager; plan de recuperación.'],['Privacidad','Consentimientos, solicitudes y retención','Configurar política Ley 21.719 con asesoría legal.'],['Respaldo','RDS, S3 versionado y exportación tenant','Probar restauración trimestral en entorno aislado.']
    ],[1550,3950,2860])
    callout(doc,'Riesgo crítico','Perder TENANT_SECRET_KEY impide descifrar configuraciones de integraciones ya creadas. Debe existir respaldo controlado y procedimiento de rotación.',ORANGE)
    add_heading(doc,'12. Monitoreo y operación')
    bullets(doc,['CloudWatch Logs: errores HTTP, solicitudes, auditoría y eventos operativos.', 'Alarmas: 5xx, tiempo de respuesta, CPU/memoria ECS, conexiones RDS, almacenamiento RDS, errores de jobs y antivirus caído.', 'Métricas: /api/metrics requiere METRICS_TOKEN.', 'Readiness: /api/ready debe validar RDS, S3 y antivirus antes de recibir tráfico.', 'Runbook: usar docs/PRODUCTION_RUNBOOK.md y docs/PRODUCTION_VALIDATION.md como procedimientos operativos.'])
    page_break(doc)
    add_heading(doc,'13. Validación antes de salida a producción')
    checks=['Todos los tests automáticos pasan: pnpm run check.', 'Validación estricta: pnpm run validate:production con secretos y servicios reales.', 'Acceso HTTPS, HSTS/cookie segura y CORS restringido al dominio productivo.', 'MFA obligatorio para administrador y cliente.', 'Empresa A, B y C no pueden consultar, modificar ni descargar archivos de otra empresa.', 'Carga de documento pasa a S3 privado, obtiene hash y resultado antivirus.', 'Restauración de backup RDS y recuperación de un respaldo por tenant probadas.', 'Flujo: cliente → contrato → orden → persona → EPP/vehículo/hotel → libro de obra → reporte.', 'Flujo de firma: solicitud, callback firmado/rechazado y auditoría.', 'Portal cliente/mandante queda limitado por scope y vencimiento.', 'Monitoreo, alertas de caída y contacto de soporte definidos.', 'Documentación y responsable de incidentes, seguridad y datos personales aprobados.']
    bullets(doc,checks)
    add_heading(doc,'14. Entregables que debe recibir el integrador')
    add_table(doc,['Ruta / activo','Propósito'],[
      ['README.md','Punto de entrada de desarrollo y operación.'],['docs/CLOUD_ARCHITECTURE.md','Modelo multiempresa, seguridad, roles e integraciones.'],['docs/AWS_PRODUCTION_READY.md','Lista de infraestructura y configuración AWS.'],['docs/AWS_RDS_DEPLOY.md','Despliegue y base de datos RDS.'],['docs/PRODUCTION_RUNBOOK.md','Operación recurrente y respuesta a incidentes.'],['docs/PRODUCTION_VALIDATION.md','Validaciones de liberación.'],['docs/ADMINISTRACION_CLIENTES_SAAS.md','Control SaaS por empresa.'],['infra/aws/','Plantillas ECS, CI/CD, entorno y guía AWS.'],['database/postgres/001-021','Esquema y migraciones ordenadas.'],['server/routes/','Contratos de API y reglas de negocio.'],['public/','Frontend productivo.'],['server/test/','Pruebas de seguridad, multiempresa, QA y funcionalidad.']
    ],[3400,4960])
    add_heading(doc,'15. Cierre y decisión de implementación')
    add_body(doc,'Nexo Klar está preparado en código para operación SaaS multiempresa. La salida real a producción requiere que el integrador complete la infraestructura externa, configure credenciales reales, aplique migraciones en un ambiente controlado y ejecute la lista de validación. No se debe declarar producción operativa sin pruebas de restauración, aislamiento, MFA, S3 privado y monitoreo activo.')
    callout(doc,'Resultado esperado','Una plataforma con información operacional centralizada y trazable, empresas separadas, usuarios controlados, evidencia privada, continuidad operacional y capacidad de crecer por módulos e integraciones.',INDIGO)
    doc.save(OUT)
    print(OUT)

if __name__=='__main__':build()
